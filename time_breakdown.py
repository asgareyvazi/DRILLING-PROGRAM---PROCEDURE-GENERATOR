# ============================================================================
# TIME BREAKDOWN MODULE FOR DRILLING PROGRAM
# File: time_breakdown.py
# Professional Time Breakdown / AFE / Operation Schedule
# Based on NIOC / IADC format
# ============================================================================

import sqlite3
import json
import os
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from pathlib import Path

from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QTableWidget, QTableWidgetItem, QGroupBox, QComboBox,
    QSplitter, QWidget, QScrollArea, QFrame, QMessageBox,
    QTabWidget, QApplication, QCheckBox, QLineEdit,
    QHeaderView, QAbstractItemView, QFileDialog, QProgressBar,
    QSpinBox, QDoubleSpinBox, QFormLayout, QTextEdit, QMenu,
    QInputDialog
)
from PySide6.QtCore import Qt, Signal, QSize
from PySide6.QtGui import QFont, QColor, QAction


# ============================================================================
# IADC OPERATION CODES
# ============================================================================

IADC_CODES = {
    "Drlg": "Drilling",
    "Trip": "Tripping",
    "Circulation": "Circulating",
    "CMT": "Cementing",
    "CSG": "Casing/Liner Running",
    "Log": "Logging/Survey",
    "BOP Test": "BOP Testing",
    "Milling": "Milling/Reaming",
    "Completion": "Completion",
    "Dry Test": "Dry Test/Formation Test",
    "Fishing": "Fishing",
    "WOC": "Wait on Cement",
    "Rig Service": "Rig Service/Repair",
    "Cut & Slip": "Cut & Slip Drilling Line",
    "DST": "Drill Stem Test",
    "Perf": "Perforation",
    "Stimulation": "Stimulation/Acidizing",
    "Rig Move": "Rig Move",
    "Other": "Other Operations",
    "NPT": "Non-Productive Time",
    "Weather": "Weather Downtime",
}


# ============================================================================
# DATA MODEL
# ============================================================================

@dataclass
class TimeBreakdownRow:
    """یک ردیف از جدول Time Breakdown"""
    id: int = 0
    project_id: int = 0
    row_number: int = 0
    section_name: str = ""  # e.g., "8-1/2\" Hole Section"
    operation_description: str = ""
    iadc_code: str = ""
    interval_m: float = 0.0  # meters drilled
    formation: str = ""
    depth_m: float = 0.0  # depth at start
    rop: float = 0.0  # m/hr or ft/hr
    duration_hours: float = 0.0
    duration_days: float = 0.0
    cumulative_days: float = 0.0
    is_section_header: bool = False
    is_contingency: bool = False
    is_rig_less: bool = False
    remarks: str = ""


@dataclass
class TimeBreakdownProject:
    """پروژه Time Breakdown"""
    id: int = 0
    name: str = ""
    well_name: str = ""
    field_name: str = ""
    operator: str = ""
    unit_system: str = "Metric"  # Metric or Imperial
    rows: List[TimeBreakdownRow] = field(default_factory=list)
    contingency_rows: List[TimeBreakdownRow] = field(default_factory=list)
    notes: str = ""
    created_date: str = ""
    modified_date: str = ""


# ============================================================================
# DATABASE
# ============================================================================

class TimeBreakdownDatabase:
    """دیتابیس Time Breakdown"""

    def __init__(self, db_path: str = None):
        if db_path is None:
            app_dir = Path.home() / ".drilling_program"
            app_dir.mkdir(exist_ok=True)
            db_path = str(app_dir / "time_breakdown.db")

        self.db_path = db_path
        self.conn = sqlite3.connect(db_path)
        self.conn.row_factory = sqlite3.Row
        self._create_tables()

    def _create_tables(self):
        self.conn.executescript("""
            CREATE TABLE IF NOT EXISTS tb_projects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                well_name TEXT DEFAULT '',
                field_name TEXT DEFAULT '',
                operator TEXT DEFAULT '',
                unit_system TEXT DEFAULT 'Metric',
                notes TEXT DEFAULT '',
                created_date TEXT,
                modified_date TEXT
            );

            CREATE TABLE IF NOT EXISTS tb_rows (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER NOT NULL,
                row_number INTEGER NOT NULL,
                section_name TEXT DEFAULT '',
                operation_description TEXT NOT NULL,
                iadc_code TEXT DEFAULT '',
                interval_m REAL DEFAULT 0,
                formation TEXT DEFAULT '',
                depth_m REAL DEFAULT 0,
                rop REAL DEFAULT 0,
                duration_hours REAL DEFAULT 0,
                duration_days REAL DEFAULT 0,
                cumulative_days REAL DEFAULT 0,
                is_section_header INTEGER DEFAULT 0,
                is_contingency INTEGER DEFAULT 0,
                is_rig_less INTEGER DEFAULT 0,
                remarks TEXT DEFAULT '',
                FOREIGN KEY (project_id) REFERENCES tb_projects(id)
                    ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_tb_rows_proj
                ON tb_rows(project_id);
        """)
        self.conn.commit()

    def close(self):
        if self.conn:
            self.conn.close()

    # ---- Projects ----

    def save_project(self, proj: TimeBreakdownProject) -> int:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        if proj.id > 0:
            self.conn.execute(
                "UPDATE tb_projects SET name=?, well_name=?, "
                "field_name=?, operator=?, unit_system=?, "
                "notes=?, modified_date=? WHERE id=?",
                (proj.name, proj.well_name, proj.field_name,
                 proj.operator, proj.unit_system, proj.notes,
                 now, proj.id))
            pid = proj.id
        else:
            cur = self.conn.execute(
                "INSERT INTO tb_projects "
                "(name, well_name, field_name, operator, unit_system, "
                "notes, created_date, modified_date) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (proj.name, proj.well_name, proj.field_name,
                 proj.operator, proj.unit_system, proj.notes,
                 now, now))
            pid = cur.lastrowid

        # Save rows
        self.conn.execute(
            "DELETE FROM tb_rows WHERE project_id=?", (pid,))

        all_rows = proj.rows + proj.contingency_rows
        for idx, row in enumerate(all_rows):
            self.conn.execute(
                "INSERT INTO tb_rows "
                "(project_id, row_number, section_name, "
                "operation_description, iadc_code, interval_m, "
                "formation, depth_m, rop, duration_hours, "
                "duration_days, cumulative_days, is_section_header, "
                "is_contingency, is_rig_less, remarks) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (pid, idx + 1, row.section_name,
                 row.operation_description, row.iadc_code,
                 row.interval_m, row.formation, row.depth_m,
                 row.rop, row.duration_hours, row.duration_days,
                 row.cumulative_days, int(row.is_section_header),
                 int(row.is_contingency), int(row.is_rig_less),
                 row.remarks))

        self.conn.commit()
        return pid

    def load_project(self, project_id: int) -> Optional[TimeBreakdownProject]:
        row = self.conn.execute(
            "SELECT * FROM tb_projects WHERE id=?",
            (project_id,)).fetchone()
        if not row:
            return None

        proj = TimeBreakdownProject(
            id=row['id'], name=row['name'],
            well_name=row['well_name'],
            field_name=row['field_name'],
            operator=row['operator'],
            unit_system=row['unit_system'],
            notes=row['notes'] or '',
            created_date=row['created_date'],
            modified_date=row['modified_date'])

        rows = self.conn.execute(
            "SELECT * FROM tb_rows WHERE project_id=? "
            "ORDER BY row_number", (project_id,)).fetchall()

        for r in rows:
            tbr = TimeBreakdownRow(
                id=r['id'], project_id=r['project_id'],
                row_number=r['row_number'],
                section_name=r['section_name'],
                operation_description=r['operation_description'],
                iadc_code=r['iadc_code'],
                interval_m=r['interval_m'],
                formation=r['formation'],
                depth_m=r['depth_m'], rop=r['rop'],
                duration_hours=r['duration_hours'],
                duration_days=r['duration_days'],
                cumulative_days=r['cumulative_days'],
                is_section_header=bool(r['is_section_header']),
                is_contingency=bool(r['is_contingency']),
                is_rig_less=bool(r['is_rig_less']),
                remarks=r['remarks'] or '')

            if tbr.is_contingency:
                proj.contingency_rows.append(tbr)
            else:
                proj.rows.append(tbr)

        return proj

    def get_all_projects(self) -> List[Dict]:
        rows = self.conn.execute(
            "SELECT id, name, well_name, operator, modified_date "
            "FROM tb_projects ORDER BY modified_date DESC"
        ).fetchall()
        return [dict(r) for r in rows]

    def delete_project(self, project_id: int):
        self.conn.execute(
            "DELETE FROM tb_projects WHERE id=?", (project_id,))
        self.conn.commit()

    def export_project(self, project_id: int, file_path: str):
        proj = self.load_project(project_id)
        if not proj:
            return

        data = {
            'name': proj.name,
            'well_name': proj.well_name,
            'field_name': proj.field_name,
            'operator': proj.operator,
            'unit_system': proj.unit_system,
            'notes': proj.notes,
            'rows': [
                {'section': r.section_name,
                 'operation': r.operation_description,
                 'iadc': r.iadc_code,
                 'interval': r.interval_m,
                 'formation': r.formation,
                 'depth': r.depth_m,
                 'rop': r.rop,
                 'hours': r.duration_hours,
                 'days': r.duration_days,
                 'cum_days': r.cumulative_days,
                 'is_header': r.is_section_header,
                 'is_contingency': r.is_contingency,
                 'is_rig_less': r.is_rig_less,
                 'remarks': r.remarks}
                for r in (proj.rows + proj.contingency_rows)
            ]
        }

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)


# ============================================================================
# STYLE
# ============================================================================

TB_STYLE = """
QDialog, QWidget { background: #0a0f1a; color: #e0e0e0;
    font-family: 'Segoe UI'; font-size: 11px; }
QLabel { color: #c0ccd8; }
QLabel#title { color: #e94560; font-size: 18px; font-weight: bold; }
QLabel#sub { color: #8899aa; font-size: 10px; }
QLabel#total { color: #27ae60; font-size: 14px; font-weight: bold; }

QTableWidget { background: #0d1525; border: 2px solid #1a2744;
    border-radius: 4px; color: #d0d8e0; gridline-color: #1a2744; }
QTableWidget::item { padding: 3px; }
QTableWidget::item:selected { background: #1a2744; color: #fff; }
QHeaderView::section { background: #0f3460; color: #e94560;
    padding: 5px 3px; border: 1px solid #0a0f1a;
    font-weight: bold; font-size: 9px; min-height: 32px; }

QGroupBox { border: 2px solid #1a2744; border-radius: 6px;
    margin-top: 10px; padding-top: 16px;
    font-weight: bold; color: #e94560; background: #0d1525; }
QGroupBox::title { subcontrol-origin: margin; left: 10px; }

QLineEdit, QComboBox, QSpinBox, QDoubleSpinBox, QTextEdit {
    background: #0d1525; border: 1px solid #1a2744;
    border-radius: 4px; padding: 4px; color: #e0e0e0; min-height: 22px; }
QComboBox QAbstractItemView { background: #0d1525; color: #e0e0e0;
    selection-background-color: #e94560; }

QTabWidget::pane { border: 1px solid #1a2744; background: #0d1525; }
QTabBar::tab { background: #0a0f1a; color: #8899aa;
    padding: 6px 12px; font-size: 10px; }
QTabBar::tab:selected { background: #1a2744; color: #e94560;
    border-bottom: 2px solid #e94560; }

QPushButton { background: #0f3460; color: #fff; border: none;
    border-radius: 5px; padding: 6px 14px; font-weight: bold;
    font-size: 10px; }
QPushButton:hover { background: #1a5276; }
QPushButton#add { background: #27ae60; }
QPushButton#del { background: #c0392b; }
QPushButton#edit { background: #e67e22; }
QPushButton#gen { background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
    stop:0 #e94560, stop:1 #c0392b);
    font-size: 12px; padding: 10px 28px; min-height: 36px; }

QProgressBar { border: 1px solid #1a2744; border-radius: 4px;
    text-align: center; color: #fff; background: #0a0f1a; }
QProgressBar::chunk { background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
    stop:0 #e94560, stop:1 #0f3460); border-radius: 3px; }
"""


# ============================================================================
# MAIN DIALOG
# ============================================================================

class TimeBreakdownDialog(QDialog):
    """دیالوگ اصلی Time Breakdown"""

    COLUMNS = [
        "Operation Description",
        "IADC Code",
        "Interval\n(m)",
        "Formation",
        "Depth\n(m)",
        "ROP",
        "Hours",
        "Days",
        "Cum.\nDays",
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("⏱️ Time Breakdown / AFE Schedule")
        self.setMinimumSize(1350, 800)
        self.setStyleSheet(TB_STYLE)

        self.db = TimeBreakdownDatabase()
        self.current_project_id = 0

        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(6)
        layout.setContentsMargins(8, 8, 8, 8)

        # Title bar
        hdr = QHBoxLayout()
        t = QLabel("⏱️ Time Breakdown / AFE Schedule")
        t.setObjectName("title")
        hdr.addWidget(t)
        hdr.addStretch()
        self.total_label = QLabel("Total: 0.0 days")
        self.total_label.setObjectName("total")
        hdr.addWidget(self.total_label)
        layout.addLayout(hdr)

        # Project info
        grp_info = QGroupBox("📋 Project Info")
        info_layout = QHBoxLayout()

        form1 = QFormLayout()
        self.txt_well = QLineEdit()
        self.txt_well.setPlaceholderText("Well Name")
        form1.addRow("Well:", self.txt_well)
        self.txt_field = QLineEdit()
        self.txt_field.setPlaceholderText("Field Name")
        form1.addRow("Field:", self.txt_field)
        info_layout.addLayout(form1)

        form2 = QFormLayout()
        self.txt_operator = QLineEdit()
        self.txt_operator.setPlaceholderText("Operator")
        form2.addRow("Operator:", self.txt_operator)
        self.cmb_unit = QComboBox()
        self.cmb_unit.addItems(["Metric (m)", "Imperial (ft)"])
        form2.addRow("Units:", self.cmb_unit)
        info_layout.addLayout(form2)

        # DB buttons
        db_btns = QVBoxLayout()
        btn_save = QPushButton("💾 Save")
        btn_save.clicked.connect(self._save_project)
        db_btns.addWidget(btn_save)
        btn_load = QPushButton("📂 Load")
        btn_load.clicked.connect(self._load_project)
        db_btns.addWidget(btn_load)
        btn_new = QPushButton("🆕 New")
        btn_new.clicked.connect(self._new_project)
        db_btns.addWidget(btn_new)
        info_layout.addLayout(db_btns)

        grp_info.setLayout(info_layout)
        layout.addWidget(grp_info)

        # Tabs: Main + Contingency
        self.tabs = QTabWidget()

        # ---- MAIN TABLE ----
        main_widget = QWidget()
        ml = QVBoxLayout(main_widget)
        ml.setContentsMargins(0, 0, 0, 0)

        # Toolbar
        tb = QHBoxLayout()
        btn_add = QPushButton("➕ Add Row")
        btn_add.setObjectName("add")
        btn_add.clicked.connect(lambda: self._add_row(self.main_table))
        tb.addWidget(btn_add)

        btn_add_hdr = QPushButton("📌 Add Section Header")
        btn_add_hdr.clicked.connect(
            lambda: self._add_section_header(self.main_table))
        tb.addWidget(btn_add_hdr)

        btn_del = QPushButton("🗑️ Delete Row")
        btn_del.setObjectName("del")
        btn_del.clicked.connect(lambda: self._del_row(self.main_table))
        tb.addWidget(btn_del)

        btn_dup = QPushButton("📋 Duplicate")
        btn_dup.clicked.connect(lambda: self._dup_row(self.main_table))
        tb.addWidget(btn_dup)

        btn_up = QPushButton("⬆️")
        btn_up.clicked.connect(lambda: self._move_row(self.main_table, -1))
        tb.addWidget(btn_up)

        btn_down = QPushButton("⬇️")
        btn_down.clicked.connect(lambda: self._move_row(self.main_table, 1))
        tb.addWidget(btn_down)

        tb.addStretch()

        btn_recalc = QPushButton("🔄 Recalculate")
        btn_recalc.clicked.connect(
            lambda: self._recalculate(self.main_table))
        tb.addWidget(btn_recalc)

        ml.addLayout(tb)

        # Table
        self.main_table = QTableWidget()
        self._setup_table(self.main_table)
        ml.addWidget(self.main_table)

        self.tabs.addTab(main_widget, "📋 Main Operations")

        # ---- CONTINGENCY TABLE ----
        cont_widget = QWidget()
        cl = QVBoxLayout(cont_widget)
        cl.setContentsMargins(0, 0, 0, 0)

        ctb = QHBoxLayout()
        btn_cadd = QPushButton("➕ Add Row")
        btn_cadd.setObjectName("add")
        btn_cadd.clicked.connect(
            lambda: self._add_row(self.cont_table))
        ctb.addWidget(btn_cadd)

        btn_cdel = QPushButton("🗑️ Delete Row")
        btn_cdel.setObjectName("del")
        btn_cdel.clicked.connect(
            lambda: self._del_row(self.cont_table))
        ctb.addWidget(btn_cdel)

        ctb.addStretch()
        btn_crecalc = QPushButton("🔄 Recalculate")
        btn_crecalc.clicked.connect(
            lambda: self._recalculate(self.cont_table))
        ctb.addWidget(btn_crecalc)
        cl.addLayout(ctb)

        self.cont_table = QTableWidget()
        self._setup_table(self.cont_table)
        cl.addWidget(self.cont_table)

        self.tabs.addTab(cont_widget, "⚠️ Contingency Plans")

        layout.addWidget(self.tabs)

        # Bottom
        bottom = QHBoxLayout()
        bottom.addStretch()

        self.progress = QProgressBar()
        self.progress.setMinimumWidth(200)
        self.progress.setVisible(False)
        bottom.addWidget(self.progress)

        btn_export_json = QPushButton("📤 Export JSON")
        btn_export_json.clicked.connect(self._export_json)
        bottom.addWidget(btn_export_json)

        btn_gen = QPushButton("📄 Generate Word Document")
        btn_gen.setObjectName("gen")
        btn_gen.clicked.connect(self._generate_word)
        bottom.addWidget(btn_gen)

        bottom.addStretch()
        layout.addLayout(bottom)

    def _setup_table(self, table: QTableWidget):
        table.setColumnCount(len(self.COLUMNS))
        table.setHorizontalHeaderLabels(self.COLUMNS)
        table.setSelectionBehavior(QAbstractItemView.SelectRows)
        table.setAlternatingRowColors(True)
        table.setMinimumHeight(300)

        # Column widths
        header = table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        for i in range(1, len(self.COLUMNS)):
            header.setSectionResizeMode(i, QHeaderView.Interactive)

        header.setMinimumSectionSize(60)
        header.setDefaultSectionSize(80)
        header.setFixedHeight(40)
        header.setDefaultAlignment(Qt.AlignCenter | Qt.TextWordWrap)

        table.verticalHeader().setDefaultSectionSize(30)

        # Context menu
        table.setContextMenuPolicy(Qt.CustomContextMenu)
        table.customContextMenuRequested.connect(
            lambda pos, t=table: self._table_context(t, pos))

    # ================================================================
    # ROW OPERATIONS
    # ================================================================

    def _add_row(self, table: QTableWidget, data: list = None):
        row = table.rowCount()
        table.insertRow(row)

        if data:
            for col, val in enumerate(data):
                if col < len(self.COLUMNS):
                    item = QTableWidgetItem(str(val))
                    if col >= 2:  # numeric columns center
                        item.setTextAlignment(Qt.AlignCenter)
                    table.setItem(row, col, item)
        else:
            # Default empty row with IADC combo
            for col in range(len(self.COLUMNS)):
                item = QTableWidgetItem("")
                if col >= 2:
                    item.setTextAlignment(Qt.AlignCenter)
                table.setItem(row, col, item)

        self._update_total()

    def _add_section_header(self, table: QTableWidget):
        name, ok = QInputDialog.getText(
            self, "Section Header",
            "Section Name (e.g., 8-1/2\" Hole Section):")
        if ok and name:
            row = table.rowCount()
            table.insertRow(row)

            item = QTableWidgetItem(name)
            item.setBackground(QColor(0x0F, 0x34, 0x60))
            item.setForeground(QColor(0xFF, 0xFF, 0xFF))
            f = item.font()
            f.setBold(True)
            item.setFont(f)
            table.setItem(row, 0, item)

            # Mark other cells
            for col in range(1, len(self.COLUMNS)):
                ci = QTableWidgetItem("")
                ci.setBackground(QColor(0x0F, 0x34, 0x60))
                table.setItem(row, col, ci)

    def _del_row(self, table: QTableWidget):
        rows = set()
        for item in table.selectedItems():
            rows.add(item.row())
        for r in sorted(rows, reverse=True):
            table.removeRow(r)
        self._update_total()

    def _dup_row(self, table: QTableWidget):
        row = table.currentRow()
        if row < 0:
            return
        data = []
        for col in range(table.columnCount()):
            item = table.item(row, col)
            data.append(item.text() if item else "")
        self._add_row(table, data)

    def _move_row(self, table: QTableWidget, direction: int):
        row = table.currentRow()
        if row < 0:
            return
        new_row = row + direction
        if new_row < 0 or new_row >= table.rowCount():
            return

        # Swap
        for col in range(table.columnCount()):
            item1 = table.takeItem(row, col)
            item2 = table.takeItem(new_row, col)
            table.setItem(row, col, item2)
            table.setItem(new_row, col, item1)

        table.setCurrentCell(new_row, 0)

    def _recalculate(self, table: QTableWidget):
        """محاسبه مجدد Days و Cumulative"""
        cum = 0.0
        for row in range(table.rowCount()):
            hours_item = table.item(row, 6)  # Hours column
            if not hours_item:
                continue

            hours_text = hours_item.text().strip()
            if not hours_text or hours_text.lower() == 'rig less':
                # Rig Less
                days_item = table.item(row, 7)
                if days_item:
                    days_item.setText("0.00")

                cum_item = table.item(row, 8)
                if not cum_item:
                    cum_item = QTableWidgetItem("")
                    cum_item.setTextAlignment(Qt.AlignCenter)
                    table.setItem(row, 8, cum_item)
                cum_item.setText(f"{cum:.1f}")
                continue

            try:
                hours = float(hours_text)
            except ValueError:
                continue

            days = hours / 24.0

            # Set days
            days_item = table.item(row, 7)
            if not days_item:
                days_item = QTableWidgetItem("")
                days_item.setTextAlignment(Qt.AlignCenter)
                table.setItem(row, 7, days_item)
            days_item.setText(f"{days:.2f}")

            # Check if section header (skip cumulative)
            desc_item = table.item(row, 0)
            if desc_item and desc_item.background().color() == QColor(0x0F, 0x34, 0x60):
                continue

            cum += days

            # Set cumulative
            cum_item = table.item(row, 8)
            if not cum_item:
                cum_item = QTableWidgetItem("")
                cum_item.setTextAlignment(Qt.AlignCenter)
                table.setItem(row, 8, cum_item)
            cum_item.setText(f"{cum:.1f}")

        self._update_total()

    def _update_total(self):
        # Get max cumulative from main table
        total = 0.0
        for row in range(self.main_table.rowCount()):
            cum_item = self.main_table.item(row, 8)
            if cum_item:
                try:
                    val = float(cum_item.text())
                    total = max(total, val)
                except ValueError:
                    pass
        self.total_label.setText(f"Total: {total:.1f} days")

    def _table_context(self, table, pos):
        menu = QMenu(self)
        menu.setStyleSheet(
            "QMenu { background: #0d1525; color: #e0e0e0; "
            "border: 1px solid #1a2744; } "
            "QMenu::item:selected { background: #e94560; }")

        menu.addAction("➕ Add Row", lambda: self._add_row(table))
        menu.addAction("📌 Add Section Header",
                        lambda: self._add_section_header(table))
        menu.addAction("📋 Duplicate Row",
                        lambda: self._dup_row(table))
        menu.addSeparator()
        menu.addAction("⬆️ Move Up",
                        lambda: self._move_row(table, -1))
        menu.addAction("⬇️ Move Down",
                        lambda: self._move_row(table, 1))
        menu.addSeparator()
        menu.addAction("🗑️ Delete Row",
                        lambda: self._del_row(table))
        menu.addSeparator()
        menu.addAction("🔄 Recalculate All",
                        lambda: self._recalculate(table))

        menu.exec(table.mapToGlobal(pos))

    # ================================================================
    # DATABASE OPERATIONS
    # ================================================================

    def _collect_project(self) -> TimeBreakdownProject:
        proj = TimeBreakdownProject()
        proj.id = self.current_project_id
        proj.well_name = self.txt_well.text()
        proj.field_name = self.txt_field.text()
        proj.operator = self.txt_operator.text()
        proj.unit_system = self.cmb_unit.currentText()
        proj.name = proj.well_name or "Untitled"

        # Main rows
        for row in range(self.main_table.rowCount()):
            proj.rows.append(self._table_row_to_data(
                self.main_table, row, False))

        # Contingency rows
        for row in range(self.cont_table.rowCount()):
            proj.contingency_rows.append(self._table_row_to_data(
                self.cont_table, row, True))

        return proj

    def _table_row_to_data(self, table, row, is_cont) -> TimeBreakdownRow:
        def get_text(col):
            item = table.item(row, col)
            return item.text().strip() if item else ""

        def get_float(col):
            t = get_text(col)
            try:
                return float(t)
            except ValueError:
                return 0.0

        desc_item = table.item(row, 0)
        is_header = False
        if desc_item and desc_item.background().color() == QColor(0x0F, 0x34, 0x60):
            is_header = True

        return TimeBreakdownRow(
            operation_description=get_text(0),
            iadc_code=get_text(1),
            interval_m=get_float(2),
            formation=get_text(3),
            depth_m=get_float(4),
            rop=get_float(5),
            duration_hours=get_float(6),
            duration_days=get_float(7),
            cumulative_days=get_float(8),
            is_section_header=is_header,
            is_contingency=is_cont,
            is_rig_less=(get_text(6).lower() == 'rig less'),
        )

    def _save_project(self):
        proj = self._collect_project()
        if not proj.name or proj.name == "Untitled":
            name, ok = QInputDialog.getText(
                self, "Project Name", "Name:")
            if ok and name:
                proj.name = name
            else:
                return

        pid = self.db.save_project(proj)
        self.current_project_id = pid
        QMessageBox.information(
            self, "Saved", f"Project saved! (ID: {pid})")

    def _load_project(self):
        projects = self.db.get_all_projects()
        if not projects:
            QMessageBox.information(self, "Empty",
                                     "No saved projects.")
            return

        items = [f"{p['name']} | {p['well_name']} | "
                 f"{p['modified_date']}" for p in projects]

        item, ok = QInputDialog.getItem(
            self, "Load Project", "Select:", items, 0, False)
        if ok:
            idx = items.index(item)
            proj = self.db.load_project(projects[idx]['id'])
            if proj:
                self._populate_from_project(proj)

    def _new_project(self):
        reply = QMessageBox.question(
            self, "New",
            "Create new project? Unsaved data will be lost.")
        if reply == QMessageBox.Yes:
            self.current_project_id = 0
            self.txt_well.clear()
            self.txt_field.clear()
            self.txt_operator.clear()
            self.main_table.setRowCount(0)
            self.cont_table.setRowCount(0)
            self._update_total()

    def _populate_from_project(self, proj: TimeBreakdownProject):
        self.current_project_id = proj.id
        self.txt_well.setText(proj.well_name)
        self.txt_field.setText(proj.field_name)
        self.txt_operator.setText(proj.operator)

        idx = self.cmb_unit.findText(proj.unit_system)
        if idx >= 0:
            self.cmb_unit.setCurrentIndex(idx)

        # Main table
        self.main_table.setRowCount(0)
        for r in proj.rows:
            data = [
                r.operation_description, r.iadc_code,
                str(r.interval_m) if r.interval_m else "",
                r.formation, str(r.depth_m) if r.depth_m else "",
                str(r.rop) if r.rop else "",
                "Rig Less" if r.is_rig_less else str(r.duration_hours),
                f"{r.duration_days:.2f}",
                f"{r.cumulative_days:.1f}",
            ]

            row_idx = self.main_table.rowCount()
            self.main_table.insertRow(row_idx)

            for col, val in enumerate(data):
                item = QTableWidgetItem(str(val))
                if col >= 2:
                    item.setTextAlignment(Qt.AlignCenter)
                if r.is_section_header:
                    item.setBackground(QColor(0x0F, 0x34, 0x60))
                    item.setForeground(QColor(0xFF, 0xFF, 0xFF))
                    f = item.font()
                    f.setBold(True)
                    item.setFont(f)
                self.main_table.setItem(row_idx, col, item)

        # Contingency table
        self.cont_table.setRowCount(0)
        for r in proj.contingency_rows:
            data = [
                r.operation_description, r.iadc_code,
                str(r.interval_m) if r.interval_m else "",
                r.formation, str(r.depth_m) if r.depth_m else "",
                str(r.rop) if r.rop else "",
                str(r.duration_hours),
                f"{r.duration_days:.2f}",
                f"{r.cumulative_days:.1f}",
            ]
            self._add_row(self.cont_table, data)

        self._update_total()

    # ================================================================
    # EXPORT
    # ================================================================

    def _export_json(self):
        proj = self._collect_project()
        path, _ = QFileDialog.getSaveFileName(
            self, "Export", f"TB_{proj.well_name}.json",
            "JSON Files (*.json)")
        if path:
            self.db.save_project(proj)
            self.db.export_project(
                self.current_project_id or 1, path)
            QMessageBox.information(self, "Done", "Exported!")

    def _generate_word(self):
        proj = self._collect_project()

        path, _ = QFileDialog.getSaveFileName(
            self, "Save Time Breakdown",
            f"Time_Breakdown_{proj.well_name or 'Well'}.docx",
            "Word Documents (*.docx)")
        if not path:
            return

        self.progress.setVisible(True)
        self.progress.setValue(0)
        QApplication.processEvents()

        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            doc = Document()
            sec = doc.sections[0]
            sec.orientation = 1  # Landscape
            sec.page_width = Cm(29.7)
            sec.page_height = Cm(21.0)
            sec.top_margin = Cm(1.5)
            sec.bottom_margin = Cm(1.5)
            sec.left_margin = Cm(1.5)
            sec.right_margin = Cm(1.5)

            self.progress.setValue(10)
            QApplication.processEvents()

            # Title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'))
            r = p.add_run(
                f"  Operation TBD of Well: {proj.well_name}  ")
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            doc.add_paragraph("")

            # Headers
            headers = [
                "Operation Description", "IADC Code",
                f"Interval\n({'m' if 'Metric' in proj.unit_system else 'ft'})",
                "Formation",
                f"Depth\n({'m' if 'Metric' in proj.unit_system else 'ft'})",
                "ROP",
                "Time\n(Hours)", "Time\n(Days)",
                "Total\n(Days)"
            ]

            all_rows = proj.rows
            num_rows = len(all_rows) + 1  # +1 header

            table = doc.add_table(rows=num_rows, cols=9)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = ""
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(h)
                cr.bold = True
                cr.font.size = Pt(8)
                cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cr.font.name = 'Calibri'
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="0C2D48"/>')
                cell._tc.get_or_add_tcPr().append(shd)

            # Column widths
            widths = [4.5, 0.8, 0.6, 0.8, 0.6, 0.4, 0.6, 0.6, 0.6]
            for row in table.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)

            self.progress.setValue(30)
            QApplication.processEvents()

            # Data rows
            for idx, r in enumerate(all_rows):
                row = table.rows[idx + 1]

                values = [
                    r.operation_description,
                    r.iadc_code,
                    f"{r.interval_m:.0f}" if r.interval_m else "",
                    r.formation,
                    f"{r.depth_m:.0f}" if r.depth_m else "",
                    f"{r.rop:.0f}" if r.rop else "",
                    "Rig Less" if r.is_rig_less
                    else f"{r.duration_hours:.0f}",
                    f"{r.duration_days:.2f}",
                    f"{r.cumulative_days:.1f}",
                ]

                for i, val in enumerate(values):
                    cell = row.cells[i]
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    if i == 0:
                        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cr = cp.add_run(str(val))
                    cr.font.size = Pt(8)
                    cr.font.name = 'Calibri'

                    if r.is_section_header:
                        cr.bold = True
                        cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        shd = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="0F3460"/>')
                        cell._tc.get_or_add_tcPr().append(shd)
                    elif idx % 2 == 0:
                        shd = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
                        cell._tc.get_or_add_tcPr().append(shd)

                pct = 30 + int(50 * (idx + 1) / max(len(all_rows), 1))
                self.progress.setValue(pct)
                QApplication.processEvents()

            # Total row
            doc.add_paragraph("")
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            total_days = 0
            if proj.rows:
                for r in proj.rows:
                    total_days = max(total_days, r.cumulative_days)

            tr = tp.add_run(f"Total Duration: {total_days:.1f} days")
            tr.bold = True
            tr.font.size = Pt(14)
            tr.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

            # Contingency section
            if proj.contingency_rows:
                doc.add_page_break()

                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pPr2 = p2._p.get_or_add_pPr()
                pPr2.append(parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="C0392B" '
                    f'w:val="clear"/>'))
                r2 = p2.add_run(
                    f"  Contingency Plans: {proj.well_name}  ")
                r2.bold = True
                r2.font.size = Pt(14)
                r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                doc.add_paragraph("")

                ct = doc.add_table(
                    rows=len(proj.contingency_rows) + 1, cols=9)
                ct.style = 'Table Grid'
                ct.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, h in enumerate(headers):
                    cell = ct.rows[0].cells[i]
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(h)
                    cr.bold = True
                    cr.font.size = Pt(8)
                    cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shd = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="C0392B"/>')
                    cell._tc.get_or_add_tcPr().append(shd)

                for idx, r in enumerate(proj.contingency_rows):
                    row = ct.rows[idx + 1]
                    values = [
                        r.operation_description, r.iadc_code,
                        f"{r.interval_m:.0f}" if r.interval_m else "",
                        r.formation,
                        f"{r.depth_m:.0f}" if r.depth_m else "",
                        "", f"{r.duration_hours:.0f}",
                        f"{r.duration_days:.2f}",
                        f"{r.cumulative_days:.1f}",
                    ]
                    for i, val in enumerate(values):
                        cell = row.cells[i]
                        cell.text = ""
                        cp = cell.paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER \
                            if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                        cr = cp.add_run(str(val))
                        cr.font.size = Pt(8)

            self.progress.setValue(90)
            QApplication.processEvents()

            doc.save(path)
            self.progress.setValue(100)

            QMessageBox.information(
                self, "✅ Success",
                f"Time Breakdown generated!\n{path}")

            if os.name == 'nt':
                os.startfile(path)

        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Error",
                f"{str(e)}\n\n{traceback.format_exc()[-500:]}")
        finally:
            self.progress.setVisible(False)

    def closeEvent(self, event):
        self.db.close()
        super().closeEvent(event)