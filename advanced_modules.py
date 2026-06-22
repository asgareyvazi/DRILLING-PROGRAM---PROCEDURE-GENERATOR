# ============================================================================
# DRILLING PROGRAM & PROCEDURE GENERATOR - PROFESSIONAL EDITION
# Version 3.0
# File: advanced_modules.py (Part 4 of 5)
# Kill Sheet, Rig Floor Poster, Cost Estimator, Report Templates
# ============================================================================

import math
from datetime import datetime
from typing import List, Dict, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================================
# KILL SHEET GENERATOR
# ============================================================================

class KillSheetGenerator:
    """تولید Kill Sheet حرفه‌ای"""

    def __init__(self, project):
        self.project = project

    def generate_kill_sheet(self, doc, section_name: str = None):
        """تولید Kill Sheet برای هر سکشن"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info
        wi = self.project.well_info
        wc = self.project.well_control
        bop = self.project.bop_stack

        # Find current section data
        current_casing = None
        current_mud = None
        prev_shoe_depth = 0

        for cd in self.project.casing_design:
            if section_name and cd.section_name.lower() == section_name.lower():
                current_casing = cd
                break
            current_casing = cd

        for mp in self.project.mud_programs:
            if section_name and mp.section_name.lower() == section_name.lower():
                current_mud = mp
                break
            current_mud = mp

        # ---- KILL SHEET HEADER ----
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("WELL CONTROL KILL SHEET")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        doc.add_paragraph("")

        # ---- WELL DATA TABLE ----
        well_data_table = doc.add_table(rows=8, cols=6)
        well_data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        well_data_table.style = 'Table Grid'

        well_items = [
            ("Well Name:", ci.well_name or "", "Rig:", ci.rig_name or "",
             "Date:", datetime.now().strftime("%d-%b-%Y")),
            ("Field:", ci.field_name or "", "Operator:", ci.operator_name or "",
             "Section:", section_name or ""),
            ("TD (MD):", f"{wi.total_depth_md:,.0f} ft",
             "TD (TVD):", f"{wi.total_depth_tvd:,.0f} ft",
             "KB:", f"{wi.kb_elevation:.1f} ft"),
            ("Current Depth (MD):", "_______ ft",
             "Current Depth (TVD):", "_______ ft",
             "Hole Size:", f'{current_casing.hole_size}"' if current_casing else ''),
            ("Last Casing:", f'{current_casing.casing_od}"' if current_casing else "",
             "Shoe Depth (MD):", f"{current_casing.setting_depth_md:,.0f} ft" if current_casing else "",
             "Shoe Depth (TVD):", f"{current_casing.setting_depth_tvd:,.0f} ft" if current_casing else ""),
            ("Current MW:", f"{current_mud.mud_weight_out:.1f} ppg" if current_mud else "_____ ppg",
             "LOT/FIT EMW:", "_______ ppg",
             "MAASP:", f"{wc.maasp_surface:,.0f} psi" if wc.maasp_surface > 0 else "_____ psi"),
            ("BOP WP:", f"{bop.working_pressure:,.0f} psi",
             "Choke Line ID:", f'{bop.choke_line_size}"',
             "Kill Line ID:", f'{bop.kill_line_size}"'),
            ("Pipe Ram Size:", bop.pipe_ram_size,
             "Annular WP:", f"{bop.annular_preventer_wp:,.0f} psi",
             "Diverter:", f'{bop.diverter_size}"'),
        ]

        for row_idx, items in enumerate(well_items):
            row = well_data_table.rows[row_idx]
            for col_idx in range(0, 6, 2):
                label_idx = col_idx
                value_idx = col_idx + 1
                if label_idx < len(items):
                    cell_label = row.cells[label_idx]
                    cell_label.text = ""
                    p = cell_label.paragraphs[0]
                    r = p.add_run(items[label_idx])
                    r.bold = True
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x0C, 0x2D, 0x48)
                    r.font.name = 'Calibri'
                    TableHelper.set_cell_shading(cell_label, "EBF5FB")

                if value_idx < len(items):
                    cell_value = row.cells[value_idx]
                    cell_value.text = ""
                    p = cell_value.paragraphs[0]
                    r = p.add_run(items[value_idx])
                    r.font.size = Pt(8)
                    r.font.name = 'Calibri'

        doc.add_paragraph("")

        # ---- DRILL STRING DATA ----
        ds_header = doc.add_paragraph()
        pPr = ds_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = ds_header.add_run("  DRILL STRING DATA")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        ds_table = doc.add_table(rows=7, cols=6)
        ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ds_table.style = 'Table Grid'

        ds_headers = ["Component", "OD (in)", "ID (in)",
                      "Length (ft)", "Capacity (bbl/ft)", "Volume (bbl)"]
        for i, h in enumerate(ds_headers):
            TableHelper.format_header_cell(
                ds_table.rows[0].cells[i], h, font_size=8)

        ds_components = [
            "Drill Pipe", "HWDP", "Drill Collars",
            "MWD/LWD", "Motor/RSS", "Total DS"
        ]
        for i, comp in enumerate(ds_components):
            row = ds_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], comp, font_size=8, bold=(comp == "Total DS"),
                alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for j in range(1, 6):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=8)
            if comp == "Total DS":
                for j in range(6):
                    TableHelper.set_cell_shading(
                        row.cells[j], "FEF9E7")

        doc.add_paragraph("")

        # ---- ANNULUS DATA ----
        ann_header = doc.add_paragraph()
        pPr = ann_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = ann_header.add_run("  ANNULUS DATA")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        ann_table = doc.add_table(rows=5, cols=6)
        ann_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ann_table.style = 'Table Grid'

        ann_headers = ["Annulus Section", "Hole/Csg ID (in)",
                       "Pipe OD (in)", "Length (ft)",
                       "Capacity (bbl/ft)", "Volume (bbl)"]
        for i, h in enumerate(ann_headers):
            TableHelper.format_header_cell(
                ann_table.rows[0].cells[i], h, font_size=8)

        ann_sections = [
            "Open Hole Annulus", "Cased Hole Annulus",
            "Choke/Kill Line", "Total Annulus"
        ]
        for i, sec in enumerate(ann_sections):
            row = ann_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], sec, font_size=8,
                bold=(sec == "Total Annulus"),
                alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for j in range(1, 6):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=8)
            if sec == "Total Annulus":
                for j in range(6):
                    TableHelper.set_cell_shading(
                        row.cells[j], "FEF9E7")

        doc.add_paragraph("")

        # ---- SLOW PUMP DATA ----
        sp_header = doc.add_paragraph()
        pPr = sp_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = sp_header.add_run("  SLOW PUMP RATE DATA")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        sp_table = doc.add_table(rows=4, cols=6)
        sp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sp_table.style = 'Table Grid'

        sp_headers = ["Pump", "Rate (SPM)", "Pressure (psi)",
                      "Output (bbl/stk)", "Strokes to Bit",
                      "Strokes Bit to Choke"]
        for i, h in enumerate(sp_headers):
            TableHelper.format_header_cell(
                sp_table.rows[0].cells[i], h, font_size=8)

        pump_names = ["Pump #1", "Pump #2", "Pump #3"]
        for i, pump in enumerate(pump_names):
            row = sp_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], pump, font_size=8,
                alignment=WD_ALIGN_PARAGRAPH.LEFT)

            # Pre-fill slow pump data
            if i == 0 and wc.slow_pump_rate_1 > 0:
                TableHelper.format_data_cell(
                    row.cells[1],
                    f"{wc.slow_pump_rate_1:.0f}", font_size=8)
                TableHelper.format_data_cell(
                    row.cells[2],
                    f"{wc.slow_pump_pressure_1:,.0f}", font_size=8)
            elif i == 1 and wc.slow_pump_rate_2 > 0:
                TableHelper.format_data_cell(
                    row.cells[1],
                    f"{wc.slow_pump_rate_2:.0f}", font_size=8)
                TableHelper.format_data_cell(
                    row.cells[2],
                    f"{wc.slow_pump_pressure_2:,.0f}", font_size=8)
            else:
                for j in range(1, 6):
                    TableHelper.format_data_cell(
                        row.cells[j], "", font_size=8)

        doc.add_paragraph("")

        # ---- KICK DATA (To be filled) ----
        kick_header = doc.add_paragraph()
        pPr = kick_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="C0392B" w:val="clear"/>'
        )
        pPr.append(shading)
        run = kick_header.add_run("  KICK DATA (FILL IN WHEN KICK OCCURS)")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        kick_table = doc.add_table(rows=6, cols=4)
        kick_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        kick_table.style = 'Table Grid'

        kick_items = [
            ("Time of Kick:", "____________",
             "Depth at Kick (MD):", "____________ ft"),
            ("SIDPP:", "____________ psi",
             "SICP:", "____________ psi"),
            ("Pit Gain:", "____________ bbl",
             "Type of Kick:", "☐ Gas  ☐ Salt Water  ☐ Oil"),
            ("Kill Mud Weight:", "____________ ppg",
             "ICP:", "____________ psi"),
            ("FCP:", "____________ psi",
             "Kill Method:", f"☐ Driller's  ☐ W&W"),
            ("Strokes to Kill:", "____________",
             "Time to Kill:", "____________ hrs"),
        ]

        for row_idx, items in enumerate(kick_items):
            row = kick_table.rows[row_idx]
            for col_idx in range(4):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    TableHelper.set_cell_shading(cell, "FDEDEC")

        doc.add_paragraph("")

        # ---- KILL CALCULATIONS ----
        calc_header = doc.add_paragraph()
        pPr = calc_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = calc_header.add_run("  KILL CALCULATIONS")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        formulas = [
            "Kill Mud Weight (ppg) = Original MW + (SIDPP / (0.052 × TVD))",
            "ICP (psi) = Slow Pump Pressure + SIDPP",
            "FCP (psi) = Slow Pump Pressure × (Kill MW / Original MW)",
            "MAASP (psi) = (LOT EMW - Current MW) × 0.052 × Shoe TVD",
            "Kick Tolerance (bbl) = MAASP / (0.052 × Kick Intensity × Ann Cap)",
        ]

        for f in formulas:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"  ▶ {f}")
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

        doc.add_paragraph("")

        # ---- PRESSURE SCHEDULE TABLE ----
        ps_header = doc.add_paragraph()
        pPr = ps_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = ps_header.add_run(
            "  DRILL PIPE PRESSURE SCHEDULE (ICP to FCP)")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        schedule_table = doc.add_table(rows=12, cols=3)
        schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        schedule_table.style = 'Table Grid'

        sched_headers = ["Strokes", "Drill Pipe Pressure (psi)",
                         "% of Strokes to Bit"]
        for i, h in enumerate(sched_headers):
            TableHelper.format_header_cell(
                schedule_table.rows[0].cells[i], h, font_size=8)

        for i in range(11):
            row = schedule_table.rows[i + 1]
            percent = i * 10
            TableHelper.format_data_cell(
                row.cells[0], "", font_size=8)
            TableHelper.format_data_cell(
                row.cells[1], "", font_size=8)
            TableHelper.format_data_cell(
                row.cells[2], f"{percent}%", font_size=8)
            if i % 2 == 0:
                for j in range(3):
                    TableHelper.set_cell_shading(
                        row.cells[j], "EBF5FB")

        doc.add_paragraph("")

        # ---- SIGNATURES ----
        sig_table = doc.add_table(rows=2, cols=4)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.style = 'Table Grid'

        sig_headers = ["Driller", "Toolpusher",
                       "Company Man", "Date/Time"]
        for i, h in enumerate(sig_headers):
            TableHelper.format_header_cell(
                sig_table.rows[0].cells[i], h, font_size=8)
            row2_cell = sig_table.rows[1].cells[i]
            row2_cell.text = ""
            p = row2_cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(30)

        return doc


# ============================================================================
# RIG FLOOR POSTER GENERATOR
# ============================================================================

class RigFloorPosterGenerator:
    """تولید پوستر ریگ فلور"""

    def __init__(self, project):
        self.project = project

    def generate_poster(self, doc):
        """تولید پوستر"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info
        wi = self.project.well_info
        wc = self.project.well_control
        bop = self.project.bop_stack

        # Change to landscape
        section = doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # Title
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"RIG FLOOR DATA SHEET - {ci.well_name} "
            f"| {ci.operator_name} | {ci.rig_name}"
        )
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        doc.add_paragraph("")

        # Main layout table (2 columns)
        main_table = doc.add_table(rows=1, cols=2)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # LEFT COLUMN
        left_cell = main_table.rows[0].cells[0]
        left_cell.width = Inches(5.5)

        # Well Data
        self._add_poster_section(left_cell, "WELL DATA", [
            f"Well: {ci.well_name}",
            f"Field: {ci.field_name}",
            f"TD (MD/TVD): {wi.total_depth_md:,.0f} / {wi.total_depth_tvd:,.0f} ft",
            f"Well Type: {wi.well_type}",
            f"KB Elev: {wi.kb_elevation:.1f} ft",
            f"Target: {wi.target_formation}",
        ])

        # Casing Data
        casing_lines = []
        for cd in self.project.casing_design:
            casing_lines.append(
                f'{cd.casing_od}" {cd.casing_grade} @ '
                f'{cd.setting_depth_md:,.0f} ft '
                f'(Hole: {cd.hole_size}")'
            )
        self._add_poster_section(left_cell, "CASING PROGRAM", casing_lines)

        # Mud Data
        mud_lines = []
        for mp in self.project.mud_programs:
            mud_lines.append(
                f'{mp.section_name}: {mp.mud_type} | '
                f'MW: {mp.mud_weight_in}-{mp.mud_weight_out} ppg'
            )
        self._add_poster_section(left_cell, "MUD PROGRAM", mud_lines)

        # RIGHT COLUMN
        right_cell = main_table.rows[0].cells[1]
        right_cell.width = Inches(5.5)

        # BOP Data
        self._add_poster_section(right_cell, "BOP DATA", [
            f"Type: {bop.bop_type}",
            f"WP: {bop.working_pressure:,.0f} psi",
            f"Bore: {bop.bore_size}\"",
            f"Pipe Rams: {bop.pipe_ram_size}",
            f"Blind/Shear: {'Yes' if bop.shear_ram else 'No'}",
            f"Annular: {bop.annular_preventer_size}\" @ "
            f"{bop.annular_preventer_wp:,.0f} psi",
        ])

        # Well Control Data
        self._add_poster_section(right_cell, "WELL CONTROL", [
            f"Kill Method: {wc.kill_method}",
            f"MAASP: {wc.maasp_surface:,.0f} psi",
            f"Kick Tolerance: {wc.kick_tolerance:.0f} bbl",
            f"Pit Gain Alarm: {wc.pit_gain_action_level:.0f} bbl",
            f"Slow Pump #1: {wc.slow_pump_rate_1:.0f} SPM @ "
            f"{wc.slow_pump_pressure_1:,.0f} psi",
            f"Slow Pump #2: {wc.slow_pump_rate_2:.0f} SPM @ "
            f"{wc.slow_pump_pressure_2:,.0f} psi",
        ])

        # Emergency Contacts
        emergency_lines = [
            "SHUT-IN PROCEDURE: Pick up - Stop pumps - Close HCR - Close Rams",
            f"Company Man: {ci.prepared_by}",
        ]
        if wc.emergency_contacts:
            emergency_lines.append(f"Emergency: {wc.emergency_contacts[:100]}")

        self._add_poster_section(
            right_cell, "⚠ EMERGENCY", emergency_lines,
            header_color="C0392B"
        )

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Rev {ci.revision} | "
            f"Date: {datetime.now().strftime('%d-%b-%Y')} | "
            f"{ci.classification.upper()}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

        return doc

    def _add_poster_section(self, cell, title: str,
                             items: List[str],
                             header_color: str = "1B4F72"):
        """اضافه کردن بخش به پوستر"""
        # Title
        p = cell.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{header_color}" w:val="clear"/>'
        )
        pPr.append(shading)
        run = p.add_run(f"  {title}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        # Items
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(f"  • {item}")
            run.font.size = Pt(8)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Spacer
        cell.add_paragraph("").paragraph_format.space_after = Pt(4)


# ============================================================================
# DAILY DRILLING REPORT TEMPLATE
# ============================================================================

class DailyReportTemplate:
    """قالب گزارش روزانه حفاری (IADC Format)"""

    def __init__(self, project):
        self.project = project

    def generate_template(self, doc):
        """تولید قالب DDR"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DAILY DRILLING REPORT (IADC FORMAT)")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        doc.add_paragraph("")

        # Report Header Data
        header_table = doc.add_table(rows=4, cols=6)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.style = 'Table Grid'

        header_items = [
            ("Operator:", ci.operator_name, "Well:", ci.well_name,
             "Report #:", "________"),
            ("Contractor:", ci.contractor_name, "Rig:", ci.rig_name,
             "Date:", "________"),
            ("Field:", ci.field_name, "AFE #:", "",
             "Report Period:", "0600 - 0600"),
            ("Spud Date:", ci.spud_date, "Days Since Spud:", "____",
             "Planned TD:", f"{self.project.well_info.total_depth_md:,.0f} ft"),
        ]

        for row_idx, items in enumerate(header_items):
            row = header_table.rows[row_idx]
            for col_idx in range(6):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # Depth & Operations Summary
        ops_header = doc.add_paragraph()
        pPr = ops_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = ops_header.add_run("  DEPTH & OPERATIONS SUMMARY")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        depth_table = doc.add_table(rows=3, cols=6)
        depth_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        depth_table.style = 'Table Grid'

        depth_items = [
            ("Present Depth MD:", "_______ ft",
             "Present Depth TVD:", "_______ ft",
             "Hole Size:", '______"'),
            ("Depth @ 0600 Start:", "_______ ft",
             "Footage Made:", "_______ ft",
             "Average ROP:", "_______ ft/hr"),
            ("Present Activity:", "___________________________",
             "Next Activity:", "___________________________"),
        ]

        for row_idx, items in enumerate(depth_items):
            row = depth_table.rows[row_idx]
            for col_idx in range(min(6, len(items))):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # 24-Hour Operations Log
        log_header = doc.add_paragraph()
        pPr = log_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = log_header.add_run("  24-HOUR OPERATIONS LOG")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        log_headers = ["From\n(hr:min)", "To\n(hr:min)",
                       "Duration\n(hrs)", "Code",
                       "Depth From\n(ft)", "Depth To\n(ft)",
                       "Operation Description"]

        log_table = doc.add_table(rows=25, cols=7)
        log_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        log_table.style = 'Table Grid'

        for i, h in enumerate(log_headers):
            TableHelper.format_header_cell(
                log_table.rows[0].cells[i], h, font_size=7)

        for i in range(24):
            row = log_table.rows[i + 1]
            for j in range(7):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                if i % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # Mud Data
        mud_header = doc.add_paragraph()
        pPr = mud_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = mud_header.add_run("  MUD DATA")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        mud_table = doc.add_table(rows=3, cols=10)
        mud_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        mud_table.style = 'Table Grid'

        mud_headers = ["", "MW\n(ppg)", "FV\n(s/qt)",
                       "PV\n(cP)", "YP\n(lb)", "Gel\n10s/10m",
                       "FL\n(ml)", "pH", "Cl\n(ppm)",
                       "Sand\n(%)"]
        for i, h in enumerate(mud_headers):
            TableHelper.format_header_cell(
                mud_table.rows[0].cells[i], h, font_size=7)

        for i, label in enumerate(["Suction", "Returns"]):
            row = mud_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], label, font_size=7, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT)
            TableHelper.set_cell_shading(row.cells[0], "EBF5FB")
            for j in range(1, 10):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)

        doc.add_paragraph("")

        # Bit Data
        bit_header = doc.add_paragraph()
        pPr = bit_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = bit_header.add_run("  BIT DATA")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        bit_table = doc.add_table(rows=3, cols=10)
        bit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        bit_table.style = 'Table Grid'

        bit_headers = ["Bit #", "Size\n(in)", "Type",
                       "Mfr/Model", "Nozzles", "Depth In\n(ft)",
                       "Depth Out\n(ft)", "Footage\n(ft)",
                       "Hours", "IADC\nDull Grade"]
        for i, h in enumerate(bit_headers):
            TableHelper.format_header_cell(
                bit_table.rows[0].cells[i], h, font_size=7)

        for i in range(2):
            row = bit_table.rows[i + 1]
            for j in range(10):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)

        doc.add_paragraph("")

        # Time Breakdown
        time_header = doc.add_paragraph()
        pPr = time_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = time_header.add_run("  TIME BREAKDOWN (24 HRS)")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        time_table = doc.add_table(rows=3, cols=8)
        time_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        time_table.style = 'Table Grid'

        time_headers = ["Drilling\n(hrs)", "Tripping\n(hrs)",
                        "Circulating\n(hrs)", "Casing\n(hrs)",
                        "Cementing\n(hrs)", "Logging\n(hrs)",
                        "Other\n(hrs)", "NPT\n(hrs)"]
        for i, h in enumerate(time_headers):
            TableHelper.format_header_cell(
                time_table.rows[0].cells[i], h, font_size=7)

        # Current Period
        row1 = time_table.rows[1]
        TableHelper.format_data_cell(
            row1.cells[0], "This Report", font_size=7, bold=True)
        for j in range(1, 8):
            TableHelper.format_data_cell(
                row1.cells[j], "", font_size=7)

        # Cumulative
        row2 = time_table.rows[2]
        TableHelper.format_data_cell(
            row2.cells[0], "Cumulative", font_size=7, bold=True)
        TableHelper.set_cell_shading(row2.cells[0], "FEF9E7")
        for j in range(1, 8):
            TableHelper.format_data_cell(
                row2.cells[j], "", font_size=7)
            TableHelper.set_cell_shading(row2.cells[j], "FEF9E7")

        doc.add_paragraph("")

        # HSE & Remarks
        hse_header = doc.add_paragraph()
        pPr = hse_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = hse_header.add_run("  HSE & REMARKS")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        hse_table = doc.add_table(rows=4, cols=4)
        hse_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hse_table.style = 'Table Grid'

        hse_items = [
            ("LTI:", "☐ Yes  ☐ No", "First Aid:", "☐ Yes  ☐ No"),
            ("Near Miss:", "☐ Yes  ☐ No", "Environmental:", "☐ Yes  ☐ No"),
            ("Safety Meeting:", "☐ Yes", "Drill:", "☐ Fire  ☐ BOP  ☐ None"),
            ("Remarks:", "", "", ""),
        ]

        for row_idx, items in enumerate(hse_items):
            row = hse_table.rows[row_idx]
            for col_idx in range(4):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'

        # Last row - merge for remarks
        row_last = hse_table.rows[3]
        cell0 = row_last.cells[0]
        cell0.merge(row_last.cells[3])

        doc.add_paragraph("")

        # Signatures
        sig_table = doc.add_table(rows=2, cols=4)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.style = 'Table Grid'

        for i, h in enumerate(["Driller", "Toolpusher",
                                "Night Pusher", "Company Man"]):
            TableHelper.format_header_cell(
                sig_table.rows[0].cells[i], h, font_size=8)
            cell = sig_table.rows[1].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(25)

        return doc


# ============================================================================
# COST ESTIMATOR
# ============================================================================

class WellCostEstimator:
    """تخمین هزینه چاه"""

    def __init__(self, project):
        self.project = project

    def generate_cost_estimate(self, doc):
        """تولید تخمین هزینه"""
        from word_generator import DocColors, TableHelper

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("WELL COST ESTIMATE (AFE)")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'

        doc.add_paragraph("")

        # Cost Categories
        cost_headers = ["Cost Category", "Description",
                        "Estimated Cost (USD)",
                        "% of Total", "Remarks"]

        total_days = 0
        for te in self.project.time_estimates:
            total_days = max(total_days, te.cumulative_days)

        if total_days <= 0:
            total_days = 30

        # Generate estimate
        costs = [
            ["1. Rig Costs", "Rig day rate × estimated days",
             "", "", ""],
            ["   1.1 Rig Day Rate", f"{total_days:.0f} days × $____/day",
             "", "", ""],
            ["   1.2 Rig Mobilization", "Mob/Demob",
             "", "", ""],
            ["2. Drilling Services", "", "", "", ""],
            ["   2.1 Directional Drilling", "MWD/LWD/RSS/Motor",
             "", "", ""],
            ["   2.2 Mud Logging", "Geological monitoring",
             "", "", ""],
            ["   2.3 Cementing Services", "Primary & remedial",
             "", "", ""],
            ["   2.4 Wireline Logging", "Open & cased hole logs",
             "", "", ""],
            ["   2.5 Drilling Fluids", "Mud & chemicals",
             "", "", ""],
            ["   2.6 Drill Bits", "PDC & Tricone bits",
             "", "", ""],
            ["   2.7 Coring", "Conventional core barrel & analysis",
             "", "", ""],
            ["3. Tubulars", "", "", "", ""],
            ["   3.1 Casing", "All casing strings",
             "", "", ""],
            ["   3.2 Tubing", "Completion tubing",
             "", "", ""],
            ["   3.3 Wellhead & Tree", "Wellhead & Xmas tree",
             "", "", ""],
            ["4. Completion", "", "", "", ""],
            ["   4.1 Completion Equipment", "Packer, screens, etc.",
             "", "", ""],
            ["   4.2 Perforating", "TCP/Wireline perforating",
             "", "", ""],
            ["   4.3 Testing (DST)", "Formation testing",
             "", "", ""],
            ["5. Support Services", "", "", "", ""],
            ["   5.1 Transportation", "Trucking, supply boats",
             "", "", ""],
            ["   5.2 Fuel", "Diesel for rig & equipment",
             "", "", ""],
            ["   5.3 Water Supply", "Drill water & domestic",
             "", "", ""],
            ["   5.4 Waste Disposal", "Cuttings & waste management",
             "", "", ""],
            ["   5.5 Camp & Catering", "Accommodation & food",
             "", "", ""],
            ["6. Insurance & Overheads", "", "", "", ""],
            ["   6.1 Insurance", "Well control, third party",
             "", "", ""],
            ["   6.2 Supervision", "Company staff",
             "", "", ""],
            ["   6.3 Office Support", "Engineering & planning",
             "", "", ""],
            ["7. Contingency", f"10-15% of total", "", "", ""],
            ["", "", "", "", ""],
            ["TOTAL ESTIMATED WELL COST", "",
             "$_____________", "100%", ""],
        ]

        TableHelper.create_professional_table(
            doc, cost_headers, costs
        )

        doc.add_paragraph("")

        # Cost Summary
        summary = doc.add_paragraph()
        run = summary.add_run(
            f"Note: Cost estimate based on {total_days:.0f} day "
            f"drilling duration. "
            f"Actual costs may vary. "
            f"Contingency of 10-15% recommended for exploration wells, "
            f"5-10% for development wells."
        )
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

        return doc


# ============================================================================
# WELL SCHEMATIC GENERATOR (TEXT-BASED)
# ============================================================================

class WellSchematicGenerator:
    """تولید شماتیک چاه (متنی)"""

    def __init__(self, project):
        self.project = project

    def generate_schematic(self, doc):
        """تولید شماتیک ساده چاه"""
        from word_generator import DocColors, TableHelper

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("WELLBORE SCHEMATIC")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        # Create schematic table
        if not self.project.casing_design:
            doc.add_paragraph("No casing design data for schematic.")
            return doc

        schematic_table = doc.add_table(
            rows=len(self.project.casing_design) + 2, cols=5
        )
        schematic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        schematic_table.style = 'Table Grid'

        headers = ["Depth\n(ft MD)", "Hole Size\n(in)",
                    "Schematic", "Casing\n(OD × Wt × Grade)",
                    "TOC\n(ft MD)"]
        for i, h in enumerate(headers):
            TableHelper.format_header_cell(
                schematic_table.rows[0].cells[i], h, font_size=9)

        # Surface line
        row = schematic_table.rows[1]
        TableHelper.format_data_cell(row.cells[0], "0", font_size=9)
        TableHelper.format_data_cell(row.cells[1], "", font_size=9)
        TableHelper.format_data_cell(
            row.cells[2], "═══╤═══", font_size=12)
        TableHelper.format_data_cell(
            row.cells[3], "Surface / KB", font_size=9)
        TableHelper.format_data_cell(row.cells[4], "", font_size=9)

        # Casing strings
        for idx, cd in enumerate(self.project.casing_design):
            row = schematic_table.rows[idx + 2]

            TableHelper.format_data_cell(
                row.cells[0], f"{cd.setting_depth_md:,.0f}",
                font_size=9, bold=True)
            TableHelper.format_data_cell(
                row.cells[1], f'{cd.hole_size}"', font_size=9)

            # ASCII art for casing
            if idx == 0:
                symbol = "║  ╔══╗  ║"
            elif idx == len(self.project.casing_design) - 1:
                symbol = "   ║╔╗║   \n   ║╚╝║   "
            else:
                symbol = " ║ ╔══╗ ║ "

            TableHelper.format_data_cell(
                row.cells[2], symbol, font_size=10)

            casing_desc = (
                f'{cd.casing_od}" × {cd.casing_weight} ppf × '
                f'{cd.casing_grade}\n'
                f'({cd.casing_connection})'
            )
            TableHelper.format_data_cell(
                row.cells[3], casing_desc, font_size=9,
                alignment=WD_ALIGN_PARAGRAPH.LEFT)

            TableHelper.format_data_cell(
                row.cells[4], f"{cd.top_of_cement_md:,.0f}",
                font_size=9)

            # Color based on section
            colors = ["EAFAF1", "EBF5FB", "FEF9E7", "FDEDEC", "F5EEF8"]
            bg_color = colors[idx % len(colors)]
            for j in range(5):
                TableHelper.set_cell_shading(row.cells[j], bg_color)

        return doc


# ============================================================================
# TRIP SHEET TEMPLATE
# ============================================================================

class TripSheetTemplate:
    """قالب Trip Sheet"""

    def __init__(self, project):
        self.project = project

    def generate_trip_sheet(self, doc):
        """تولید Trip Sheet"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TRIP SHEET / FILL & CHECK RECORD")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        # Info
        info_table = doc.add_table(rows=2, cols=6)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Table Grid'

        info_items = [
            ("Well:", ci.well_name, "Date:", "________",
             "Trip #:", "________"),
            ("Depth (MD):", "_______ ft", "Hole Size:", '______"',
             "DP Size:", '______"'),
        ]

        for row_idx, items in enumerate(info_items):
            row = info_table.rows[row_idx]
            for col_idx in range(6):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # Trip data
        trip_headers = [
            "Stand #", "Time", "Stands\nPulled/Run",
            "Pipe Displaced\n(Calc) bbl",
            "Mud Pumped\n(Actual) bbl",
            "Difference\n(bbl)",
            "Cum.\nDifference",
            "Trip Tank\nLevel (in)",
            "Remarks"
        ]

        trip_table = doc.add_table(rows=31, cols=9)
        trip_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        trip_table.style = 'Table Grid'

        for i, h in enumerate(trip_headers):
            TableHelper.format_header_cell(
                trip_table.rows[0].cells[i], h, font_size=7)

        for i in range(30):
            row = trip_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], str(i + 1), font_size=7)
            for j in range(1, 9):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)
            if i % 2 == 0:
                for j in range(9):
                    TableHelper.set_cell_shading(
                        row.cells[j], "EBF5FB")

            # Highlight every 5th stand
            if (i + 1) % 5 == 0:
                for j in range(9):
                    TableHelper.set_cell_shading(
                        row.cells[j], "FEF9E7")

        doc.add_paragraph("")

        # Fill volume info
        p = doc.add_paragraph()
        run = p.add_run("CALCULATED FILL VOLUME: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

        run2 = p.add_run(
            "DP Displacement = _____ bbl/stand | "
            "Fill per 5 stands = _____ bbl | "
            "Trip tank sensitivity = _____ bbl/inch"
        )
        run2.font.size = Pt(9)

        doc.add_paragraph("")

        # Warning
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="FDEDEC" w:val="clear"/>'
        )
        pPr.append(shading)
        run = p.add_run(
            "  ⚠ WARNING: If difference exceeds 3 bbl (or company limit), "
            "STOP tripping and investigate. "
            "Perform flow check if pit gain observed."
        )
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        return doc


# ============================================================================
# TALLY SHEET TEMPLATE
# ============================================================================

class TallySheetTemplate:
    """قالب تالی لوله"""

    def __init__(self, project):
        self.project = project

    def generate_tally_sheet(self, doc, casing_section: str = ""):
        """تولید تالی شیت"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"CASING TALLY SHEET - {casing_section.upper()}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        # Info section
        info_table = doc.add_table(rows=2, cols=6)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Table Grid'

        info_items = [
            ("Well:", ci.well_name, "Date:", "________",
             "Section:", casing_section),
            ("Casing Size:", '______"', "Weight:", "______ ppf",
             "Grade:", "________"),
        ]

        for row_idx, items in enumerate(info_items):
            row = info_table.rows[row_idx]
            for col_idx in range(6):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # Tally table
        tally_headers = [
            "Joint #", "Joint Length\n(ft)",
            "Cumulative\nLength (ft)",
            "Thread\nCondition",
            "Centralizer\nInstalled",
            "Torque\n(ft-lbs)",
            "Fill Up\nVolume (bbl)",
            "Remarks"
        ]

        tally_table = doc.add_table(rows=41, cols=8)
        tally_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tally_table.style = 'Table Grid'

        for i, h in enumerate(tally_headers):
            TableHelper.format_header_cell(
                tally_table.rows[0].cells[i], h, font_size=7)

        for i in range(40):
            row = tally_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], str(i + 1), font_size=7)
            for j in range(1, 8):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)
            if i % 2 == 0:
                for j in range(8):
                    TableHelper.set_cell_shading(
                        row.cells[j], "EBF5FB")

            # Special rows for accessories
            if i == 0:
                TableHelper.format_data_cell(
                    row.cells[7], "Float Shoe", font_size=7,
                    color=RGBColor(0xE9, 0x45, 0x60))
                for j in range(8):
                    TableHelper.set_cell_shading(
                        row.cells[j], "FDEDEC")

        return doc


# ============================================================================
# BIT RECORD TEMPLATE
# ============================================================================

class BitRecordTemplate:
    """قالب رکورد مته"""

    def __init__(self, project):
        self.project = project

    def generate_bit_record(self, doc):
        """تولید فرم رکورد مته"""
        from word_generator import DocColors, TableHelper

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BIT RECORD (IADC DULL BIT GRADING)")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        headers = [
            "Bit #", "Size\n(in)", "Type",
            "Manufacturer\n/ Model", "Serial #",
            "Nozzles\n(32nds)", "TFA\n(sq in)",
            "Depth In\n(ft)", "Depth Out\n(ft)",
            "Footage\n(ft)", "Hours",
            "WOB\n(klbs)", "RPM",
            "Flow\n(GPM)", "ROP\n(ft/hr)",
        ]

        bit_table = doc.add_table(rows=11, cols=15)
        bit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        bit_table.style = 'Table Grid'

        for i, h in enumerate(headers):
            TableHelper.format_header_cell(
                bit_table.rows[0].cells[i], h, font_size=6)

        for i in range(10):
            row = bit_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], str(i + 1), font_size=7)
            for j in range(1, 15):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)
            if i % 2 == 0:
                for j in range(15):
                    TableHelper.set_cell_shading(
                        row.cells[j], "EBF5FB")

        doc.add_paragraph("")

        # Dull Bit Grading
        dull_header = doc.add_paragraph()
        pPr = dull_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = dull_header.add_run(
            "  IADC DULL BIT GRADING SYSTEM")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        dull_headers = [
            "Bit #", "Inner Rows\n(I)",
            "Outer Rows\n(O)", "Dull Char.\n(D)",
            "Location\n(L)", "Bearing\n(B)",
            "Gauge\n(G)", "Other Dull\n(O)",
            "Reason\nPulled"
        ]

        dull_table = doc.add_table(rows=11, cols=9)
        dull_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        dull_table.style = 'Table Grid'

        for i, h in enumerate(dull_headers):
            TableHelper.format_header_cell(
                dull_table.rows[0].cells[i], h, font_size=7)

        for i in range(10):
            row = dull_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], str(i + 1), font_size=7)
            for j in range(1, 9):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)

        # Grading Legend
        doc.add_paragraph("")
        legend_items = [
            "I/O: 0-8 (0=No wear, 8=Total loss)",
            "D: BT=Broken Teeth, CT=Chipped, ER=Erosion, "
            "FC=Flat Crested, RO=Ring Out, WT=Worn Teeth",
            "L: N=Nose, S=Shoulder, G=Gauge, A=All, M=Middle, H=Heel",
            "B: 0-8 (Fixed cutter: X=N/A)",
            "G: I=In gauge, 1/16=1/16\" under gauge",
            "Reason Pulled: BHA=BHA Change, CM=Condition Mud, "
            "CP=Core Point, DMF=Downhole Motor Failure, "
            "DP=Drill Plug, DTF=Downhole Tool Failure, "
            "FM=Formation Change, HP=Hole Problems, "
            "HR=Hours, PP=Pump Pressure, PR=Penetration Rate, "
            "RIG=Rig Repair, TD=Total Depth, TW=Twist Off",
        ]

        for item in legend_items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"  • {item}")
            run.font.size = Pt(7)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x56, 0x6C, 0x73)

        return doc


# ============================================================================
# CEMENT JOB REPORT TEMPLATE
# ============================================================================

class CementJobReportTemplate:
    """قالب گزارش سیمانکاری"""

    def __init__(self, project):
        self.project = project

    def generate_report(self, doc, section_name: str = ""):
        """تولید گزارش سیمانکاری"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"CEMENT JOB REPORT - {section_name.upper()}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        # General Info
        info_table = doc.add_table(rows=3, cols=6)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Table Grid'

        info_items = [
            ("Well:", ci.well_name, "Date:", "________",
             "Section:", section_name),
            ("Operator:", ci.operator_name,
             "Rig:", ci.rig_name,
             "Cement Co.:", "________"),
            ("Casing Size:", '______"', "Shoe Depth:", "______ ft",
             "TOC Plan:", "______ ft"),
        ]

        for row_idx, items in enumerate(info_items):
            row = info_table.rows[row_idx]
            for col_idx in range(6):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # Planned vs Actual
        comparison_headers = [
            "Parameter", "Planned", "Actual", "Variance"
        ]

        comparison_items = [
            ["Lead Slurry Density (ppg)", "", "", ""],
            ["Lead Slurry Volume (bbl)", "", "", ""],
            ["Tail Slurry Density (ppg)", "", "", ""],
            ["Tail Slurry Volume (bbl)", "", "", ""],
            ["Spacer Volume (bbl)", "", "", ""],
            ["Displacement Volume (bbl)", "", "", ""],
            ["Displacement Rate (bpm)", "", "", ""],
            ["Max Surface Pressure (psi)", "", "", ""],
            ["Plug Bump Pressure (psi)", "", "", ""],
            ["Total Cement (sacks)", "", "", ""],
            ["Excess (%)", "", "", ""],
            ["Cement Returns", "", "", ""],
            ["WOC Time (hrs)", "", "", ""],
        ]

        TableHelper.create_professional_table(
            doc, comparison_headers, comparison_items
        )

        doc.add_paragraph("")

        # Job Events Log
        events_header = doc.add_paragraph()
        pPr = events_header._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="1B4F72" w:val="clear"/>'
        )
        pPr.append(shading)
        run = events_header.add_run("  JOB EVENTS LOG")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        events_table = doc.add_table(rows=16, cols=4)
        events_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        events_table.style = 'Table Grid'

        ev_headers = ["Time", "Event", "Volume (bbl)",
                      "Pressure (psi)"]
        for i, h in enumerate(ev_headers):
            TableHelper.format_header_cell(
                events_table.rows[0].cells[i], h, font_size=8)

        for i in range(15):
            row = events_table.rows[i + 1]
            for j in range(4):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=8)
            if i % 2 == 0:
                for j in range(4):
                    TableHelper.set_cell_shading(
                        row.cells[j], "EBF5FB")

        doc.add_paragraph("")

        # Job Evaluation
        eval_items = [
            "Job Rating: ☐ Excellent  ☐ Good  ☐ Fair  ☐ Poor",
            "Cement Returns at Surface: ☐ Yes  ☐ No  ☐ Partial",
            "Plug Bumped Successfully: ☐ Yes  ☐ No",
            "CBL/CBIL Scheduled: ☐ Yes  ☐ No  Date: ________",
            "Remedial Cementing Required: ☐ Yes  ☐ No",
        ]

        for item in eval_items:
            p = doc.add_paragraph()
            run = p.add_run(f"  {item}")
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

        doc.add_paragraph("")

        # Remarks
        p = doc.add_paragraph()
        run = p.add_run("REMARKS / LESSONS LEARNED:")
        run.bold = True
        run.font.size = Pt(10)

        for i in range(5):
            doc.add_paragraph(
                "___________________________________________________"
                "___________________________________________________"
            )

        return doc


# ============================================================================
# SURVEY SHEET TEMPLATE
# ============================================================================

class SurveySheetTemplate:
    """قالب شیت سروی"""

    def __init__(self, project):
        self.project = project

    def generate_survey_sheet(self, doc):
        """تولید شیت سروی"""
        from word_generator import DocColors, TableHelper

        ci = self.project.company_info

        # Header
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DIRECTIONAL SURVEY SHEET")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        # Info
        info_table = doc.add_table(rows=2, cols=6)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Table Grid'

        info_items = [
            ("Well:", ci.well_name, "Survey Tool:", "________",
             "Section:", "________"),
            ("Run #:", "________", "Date:", "________",
             "Surveyor:", "________"),
        ]

        for row_idx, items in enumerate(info_items):
            row = info_table.rows[row_idx]
            for col_idx in range(6):
                cell = row.cells[col_idx]
                cell.text = ""
                p_cell = cell.paragraphs[0]
                r = p_cell.add_run(items[col_idx])
                r.bold = (col_idx % 2 == 0)
                r.font.size = Pt(8)
                r.font.name = 'Calibri'
                if col_idx % 2 == 0:
                    TableHelper.set_cell_shading(cell, "EBF5FB")

        doc.add_paragraph("")

        # Survey Data
        survey_headers = [
            "Stn #", "MD\n(ft)", "Inc\n(°)",
            "Azi\n(°)", "TVD\n(ft)", "N/S\n(ft)",
            "E/W\n(ft)", "VS\n(ft)",
            "DLS\n(°/100ft)", "Closure\nDist (ft)",
            "Closure\nDir (°)", "Remarks"
        ]

        survey_table = doc.add_table(rows=31, cols=12)
        survey_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        survey_table.style = 'Table Grid'

        for i, h in enumerate(survey_headers):
            TableHelper.format_header_cell(
                survey_table.rows[0].cells[i], h, font_size=7)

        for i in range(30):
            row = survey_table.rows[i + 1]
            TableHelper.format_data_cell(
                row.cells[0], str(i + 1), font_size=7)
            for j in range(1, 12):
                TableHelper.format_data_cell(
                    row.cells[j], "", font_size=7)
            if i % 2 == 0:
                for j in range(12):
                    TableHelper.set_cell_shading(
                        row.cells[j], "EBF5FB")

        return doc


# ============================================================================
# APPENDIX GENERATOR - INTEGRATES ALL TEMPLATES
# ============================================================================

class AppendixGenerator:
    """تولید تمام ضمایم"""

    def __init__(self, project):
        self.project = project

    def generate_all_appendices(self, doc, progress=None):
        """تولید تمام ضمایم در سند"""
        from word_generator import DocColors

        # Kill Sheet
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX I: KILL SHEET")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        kill_sheet = KillSheetGenerator(self.project)
        kill_sheet.generate_kill_sheet(doc)

        # Wellbore Schematic
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX A: WELLBORE SCHEMATIC")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        schematic = WellSchematicGenerator(self.project)
        schematic.generate_schematic(doc)

        # Trip Sheet
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: TRIP SHEET TEMPLATE")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        trip_sheet = TripSheetTemplate(self.project)
        trip_sheet.generate_trip_sheet(doc)

        # Bit Record
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: BIT RECORD")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        bit_record = BitRecordTemplate(self.project)
        bit_record.generate_bit_record(doc)

        # Tally Sheet
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: CASING TALLY SHEET")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        tally = TallySheetTemplate(self.project)
        for cd in self.project.casing_design:
            tally.generate_tally_sheet(doc, cd.section_name)
            doc.add_page_break()

        # Survey Sheet
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: SURVEY SHEET")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        survey = SurveySheetTemplate(self.project)
        survey.generate_survey_sheet(doc)

        # Cement Job Report
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: CEMENT JOB REPORT TEMPLATE")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        cement_report = CementJobReportTemplate(self.project)
        cement_report.generate_report(doc)

        # Cost Estimate
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: WELL COST ESTIMATE")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        cost = WellCostEstimator(self.project)
        cost.generate_cost_estimate(doc)

        # Daily Report Template
        doc.add_page_break()
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'
        )
        pPr.append(shading)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("APPENDIX: DAILY DRILLING REPORT TEMPLATE")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        ddr = DailyReportTemplate(self.project)
        ddr.generate_template(doc)

        return doc


# ============================================================================
# END OF PART 4
# ============================================================================