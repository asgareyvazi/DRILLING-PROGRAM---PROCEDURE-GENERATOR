# ============================================================================
# DRILLING PROGRAM & PROCEDURE GENERATOR - PROFESSIONAL EDITION
# Version 3.0
# File: word_generator.py (Part 3 of 5)
# Professional Word Document Generator - Section 1
# ============================================================================

import os
import math
from datetime import datetime
from typing import List, Dict, Optional, Tuple

# python-docx imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Mm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. Install with: pip install python-docx")


# ============================================================================
# COLOR SCHEME - PROFESSIONAL OIL & GAS THEME
# ============================================================================

class DocColors:
    """رنگ‌های سند حرفه‌ای"""
    # Primary Colors
    DARK_NAVY = RGBColor(0x0C, 0x2D, 0x48)      # Dark Navy Blue
    NAVY = RGBColor(0x0F, 0x34, 0x60)             # Navy Blue
    STEEL_BLUE = RGBColor(0x1B, 0x4F, 0x72)       # Steel Blue
    MEDIUM_BLUE = RGBColor(0x21, 0x61, 0x8C)      # Medium Blue
    LIGHT_BLUE = RGBColor(0x2E, 0x86, 0xC1)       # Light Blue
    SKY_BLUE = RGBColor(0xD4, 0xE6, 0xF1)         # Sky Blue (light bg)

    # Accent Colors
    RED = RGBColor(0xE9, 0x45, 0x60)               # Accent Red
    DARK_RED = RGBColor(0xC0, 0x39, 0x2B)          # Dark Red
    ORANGE = RGBColor(0xE6, 0x7E, 0x22)            # Orange
    GREEN = RGBColor(0x27, 0xAE, 0x60)             # Green
    DARK_GREEN = RGBColor(0x1E, 0x8C, 0x49)        # Dark Green
    GOLD = RGBColor(0xF3, 0x9C, 0x12)              # Gold
    YELLOW = RGBColor(0xF1, 0xC4, 0x0F)            # Yellow

    # Neutral Colors
    BLACK = RGBColor(0x17, 0x20, 0x2A)             # Near Black
    DARK_GRAY = RGBColor(0x2C, 0x3E, 0x50)         # Dark Gray
    MEDIUM_GRAY = RGBColor(0x56, 0x6C, 0x73)       # Medium Gray
    GRAY = RGBColor(0x85, 0x92, 0x9E)              # Gray
    LIGHT_GRAY = RGBColor(0xD5, 0xDB, 0xDB)        # Light Gray
    VERY_LIGHT = RGBColor(0xEB, 0xED, 0xEE)        # Very Light Gray
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)              # White

    # Table Colors (hex strings for cell shading)
    HEADER_BG = "0C2D48"
    SUBHEADER_BG = "1B4F72"
    ROW_ALT_BG = "EBF5FB"
    ROW_NORMAL_BG = "FFFFFF"
    CAUTION_BG = "FDEDEC"
    WARNING_BG = "FEF9E7"
    SUCCESS_BG = "EAFAF1"
    INFO_BG = "EBF5FB"


# ============================================================================
# DOCUMENT STYLE MANAGER
# ============================================================================

class DocumentStyleManager:
    """مدیریت استایل‌های سند"""

    def __init__(self, doc: 'Document'):
        self.doc = doc
        self.setup_styles()

    def setup_styles(self):
        """ایجاد تمام استایل‌های مورد نیاز"""
        styles = self.doc.styles

        # ---- Title Style ----
        self._create_or_modify_style(
            'DocTitle', 'Title',
            font_name='Calibri', font_size=28,
            font_color=DocColors.DARK_NAVY,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(6)
        )

        # ---- Heading 1 ----
        self._create_or_modify_style(
            'Heading 1', 'Heading 1',
            font_name='Calibri', font_size=18,
            font_color=DocColors.WHITE,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(18), space_after=Pt(8)
        )

        # ---- Heading 2 ----
        self._create_or_modify_style(
            'Heading 2', 'Heading 2',
            font_name='Calibri', font_size=14,
            font_color=DocColors.DARK_NAVY,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(14), space_after=Pt(6)
        )

        # ---- Heading 3 ----
        self._create_or_modify_style(
            'Heading 3', 'Heading 3',
            font_name='Calibri', font_size=12,
            font_color=DocColors.STEEL_BLUE,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(10), space_after=Pt(4)
        )

        # ---- Normal Body ----
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.font.color.rgb = DocColors.BLACK
        pf = normal.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.15

        # ---- Procedure Step Style ----
        try:
            proc_style = styles.add_style('ProcedureStep', WD_STYLE_TYPE.PARAGRAPH)
            proc_style.font.name = 'Calibri'
            proc_style.font.size = Pt(10)
            proc_style.font.color.rgb = DocColors.BLACK
            proc_style.paragraph_format.left_indent = Cm(0.5)
            proc_style.paragraph_format.space_after = Pt(2)
            proc_style.paragraph_format.line_spacing = 1.2
        except ValueError:
            pass

        # ---- Caption Style ----
        try:
            cap_style = styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
            cap_style.font.name = 'Calibri'
            cap_style.font.size = Pt(9)
            cap_style.font.italic = True
            cap_style.font.color.rgb = DocColors.MEDIUM_GRAY
            cap_style.paragraph_format.space_before = Pt(4)
            cap_style.paragraph_format.space_after = Pt(8)
            cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except ValueError:
            pass

        # ---- Note/Warning Style ----
        try:
            note_style = styles.add_style('NoteStyle', WD_STYLE_TYPE.PARAGRAPH)
            note_style.font.name = 'Calibri'
            note_style.font.size = Pt(9)
            note_style.font.italic = True
            note_style.font.color.rgb = DocColors.STEEL_BLUE
            note_style.paragraph_format.left_indent = Cm(1.0)
            note_style.paragraph_format.space_after = Pt(4)
        except ValueError:
            pass

        # ---- Footer Style ----
        try:
            foot_style = styles.add_style('FooterStyle', WD_STYLE_TYPE.PARAGRAPH)
            foot_style.font.name = 'Calibri'
            foot_style.font.size = Pt(8)
            foot_style.font.color.rgb = DocColors.GRAY
            foot_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except ValueError:
            pass

    def _create_or_modify_style(self, name, base_name, font_name='Calibri',
                                 font_size=11, font_color=None, bold=False,
                                 italic=False, alignment=None,
                                 space_before=None, space_after=None):
        """ایجاد یا اصلاح استایل"""
        styles = self.doc.styles
        try:
            style = styles[name]
        except KeyError:
            try:
                style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            except ValueError:
                style = styles[name]

        style.font.name = font_name
        style.font.size = Pt(font_size)
        if font_color:
            style.font.color.rgb = font_color
        style.font.bold = bold
        style.font.italic = italic
        if alignment:
            style.paragraph_format.alignment = alignment
        if space_before:
            style.paragraph_format.space_before = space_before
        if space_after:
            style.paragraph_format.space_after = space_after


# ============================================================================
# TABLE HELPER
# ============================================================================

class TableHelper:
    """ابزار ساخت جداول حرفه‌ای"""

    @staticmethod
    def set_cell_shading(cell, color_hex: str):
        """تنظیم رنگ پس‌زمینه سلول"""
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def set_cell_border(cell, **kwargs):
        """تنظیم حاشیه سلول"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    @staticmethod
    def set_cell_width(cell, width_inches: float):
        """تنظیم عرض سلول"""
        cell.width = Inches(width_inches)

    @staticmethod
    def format_header_cell(cell, text: str, font_size: int = 9):
        """فرمت سلول هدر"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.WHITE
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        TableHelper.set_cell_shading(cell, DocColors.HEADER_BG)

    @staticmethod
    def format_subheader_cell(cell, text: str, font_size: int = 9):
        """فرمت سلول ساب‌هدر"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.WHITE
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        TableHelper.set_cell_shading(cell, DocColors.SUBHEADER_BG)

    @staticmethod
    def format_data_cell(cell, text: str, font_size: int = 9,
                         bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color: RGBColor = None):
        """فرمت سلول داده"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = alignment
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = color
        else:
            run.font.color.rgb = DocColors.BLACK
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @staticmethod
    def format_label_cell(cell, text: str, font_size: int = 9):
        """فرمت سلول برچسب"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.DARK_NAVY
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        TableHelper.set_cell_shading(cell, DocColors.ROW_ALT_BG)

    @staticmethod
    def format_value_cell(cell, text: str, font_size: int = 9):
        """فرمت سلول مقدار"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.BLACK
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @staticmethod
    def set_row_height(row, height_cm: float = 0.7):
        """تنظیم ارتفاع ردیف"""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(
            f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" '
            f'w:hRule="atLeast"/>'
        )
        trPr.append(trHeight)

    @staticmethod
    def create_professional_table(doc, headers: List[str],
                                   data: List[List[str]],
                                   col_widths: List[float] = None,
                                   caption: str = None,
                                   alt_row_shading: bool = True) -> 'Table':
        """ایجاد جدول حرفه‌ای کامل"""
        num_cols = len(headers)
        num_rows = len(data) + 1  # +1 for header

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Set column widths
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        # Format header row
        header_row = table.rows[0]
        TableHelper.set_row_height(header_row, 0.8)
        for i, header_text in enumerate(headers):
            TableHelper.format_header_cell(header_row.cells[i], header_text)

        # Format data rows
        for row_idx, row_data in enumerate(data):
            row = table.rows[row_idx + 1]
            TableHelper.set_row_height(row, 0.6)
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    TableHelper.format_data_cell(cell, str(cell_text))
                    if alt_row_shading and row_idx % 2 == 0:
                        TableHelper.set_cell_shading(
                            cell, DocColors.ROW_ALT_BG)

        # Add caption
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_run.font.color.rgb = DocColors.MEDIUM_GRAY

        return table

    @staticmethod
    def create_key_value_table(doc, data: List[Tuple[str, str]],
                                title: str = None,
                                num_cols: int = 2) -> 'Table':
        """ایجاد جدول کلید-مقدار (مثل فرم)"""
        if not data:
            return None

        if num_cols == 4:
            num_rows = math.ceil(len(data) / 2)
            table = doc.add_table(rows=num_rows, cols=4)
        else:
            num_rows = len(data)
            table = doc.add_table(rows=num_rows, cols=2)

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        if num_cols == 4:
            for row_idx in range(num_rows):
                row = table.rows[row_idx]
                TableHelper.set_row_height(row, 0.6)

                # Left pair
                left_idx = row_idx
                if left_idx < len(data):
                    item = data[left_idx]
                    key = item[0] if len(item) > 0 else ""
                    val = item[1] if len(item) > 1 else ""
                    TableHelper.format_label_cell(row.cells[0], str(key))
                    TableHelper.format_value_cell(row.cells[1], str(val))

                # Right pair
                right_idx = row_idx + num_rows
                if right_idx < len(data):
                    item = data[right_idx]
                    key = item[0] if len(item) > 0 else ""
                    val = item[1] if len(item) > 1 else ""
                    TableHelper.format_label_cell(row.cells[2], str(key))
                    TableHelper.format_value_cell(row.cells[3], str(val))
                else:
                    # Empty right side
                    TableHelper.format_label_cell(row.cells[2], "")
                    TableHelper.format_value_cell(row.cells[3], "")
        else:
            for row_idx, item in enumerate(data):
                row = table.rows[row_idx]
                TableHelper.set_row_height(row, 0.6)
                key = item[0] if len(item) > 0 else ""
                val = item[1] if len(item) > 1 else ""
                TableHelper.format_label_cell(row.cells[0], str(key))
                TableHelper.format_value_cell(row.cells[1], str(val))

        return table

    @staticmethod
    def merge_cells_in_row(table, row_idx: int, start_col: int,
                            end_col: int):
        """ادغام سلول‌ها در یک ردیف"""
        row = table.rows[row_idx]
        cell = row.cells[start_col]
        for col in range(start_col + 1, end_col + 1):
            cell = cell.merge(row.cells[col])
        return cell


# ============================================================================
# MAIN WORD DOCUMENT GENERATOR
# ============================================================================

class DrillingProgramWordGenerator:
    """تولیدکننده اصلی سند Word"""

    def __init__(self, project):
        self.project = project
        self.doc = Document()
        self.style_mgr = DocumentStyleManager(self.doc)
        self.table_helper = TableHelper()
        self.section_counter = 0
        self.table_counter = 0
        self.figure_counter = 0
        self.page_count = 0

        # Import procedure generator
        from engineering_calculations import ProcedureGenerator, CalculationEngine
        self.proc_gen = ProcedureGenerator(project)
        self.calc_engine = CalculationEngine(project)

    def generate(self, file_path: str, progress=None):
        """تولید سند کامل"""
        if not HAS_DOCX:
            raise ImportError("python-docx is not installed")

        self._setup_document()
        self._update_progress(progress, 5)

        # 1. Cover Page
        self._create_cover_page()
        self._update_progress(progress, 10)

        # 2. Revision History
        self._create_revision_history()
        self._update_progress(progress, 12)

        # 3. Table of Contents
        self._create_table_of_contents()
        self._update_progress(progress, 14)

        # 4. Abbreviations & Definitions
        self._create_abbreviations()
        self._update_progress(progress, 16)

        # 5. Executive Summary
        self._create_executive_summary()
        self._update_progress(progress, 18)

        # 6. Well Information
        self._create_well_information_section()
        self._update_progress(progress, 25)

        # 7. Rig Specifications
        self._create_rig_specification_section()
        self._update_progress(progress, 28)

        # 8. Formation Prognosis
        self._create_formation_prognosis_section()
        self._update_progress(progress, 32)

        # 9. Hazard Analysis
        self._create_hazard_analysis_section()
        self._update_progress(progress, 36)

        # 10. Casing Design
        self._create_casing_design_section()
        self._update_progress(progress, 42)

        # 11. Drilling Fluid Program
        self._create_mud_program_section()
        self._update_progress(progress, 48)

        # 12. BHA & Drill String
        self._create_bha_section()
        self._update_progress(progress, 52)

        # 13. Hydraulics
        self._create_hydraulics_section()
        self._update_progress(progress, 55)

        # 14. Cementing Program
        self._create_cement_section()
        self._update_progress(progress, 60)

        # 15. Directional Drilling Plan
        self._create_directional_section()
        self._update_progress(progress, 64)

        # 16. BOP & Well Control
        self._create_bop_well_control_section()
        self._update_progress(progress, 68)

        # 17. Evaluation Program
        self._create_evaluation_section()
        self._update_progress(progress, 72)

        # 18. Time vs Depth Estimate
        self._create_time_estimate_section()
        self._update_progress(progress, 76)

        # 19. Detailed Procedures
        self._create_all_procedures_section()
        self._update_progress(progress, 90)

        # 20. Emergency Procedures
        self._create_emergency_section()
        self._update_progress(progress, 94)

        # 21. Appendices
        self._create_appendices()
        self._update_progress(progress, 97)

        # Add headers and footers
        self._add_headers_footers()

        # Save document
        self.doc.save(file_path)
        self._update_progress(progress, 100)

    def _update_progress(self, progress, value):
        """بروزرسانی نوار پیشرفت"""
        if progress:
            progress.setValue(value)
            from PySide6.QtWidgets import QApplication
            QApplication.processEvents()

    # ========================================================================
    # DOCUMENT SETUP
    # ========================================================================

    def _setup_document(self):
        """تنظیمات اولیه سند"""
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    # ========================================================================
    # COVER PAGE
    # ========================================================================

    def _create_cover_page(self):
        """ایجاد صفحه جلد حرفه‌ای"""
        ci = self.project.company_info
        wi = self.project.well_info

        # Top border line
        self._add_colored_line(DocColors.DARK_NAVY, width=3)

        # Spacer
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        # Classification Banner
        classification_para = self.doc.add_paragraph()
        classification_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classification_para.add_run(f"⬛ {ci.classification.upper()} ⬛")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = DocColors.RED

        self.doc.add_paragraph("")

        # Operator Name
        if ci.operator_name:
            op_para = self.doc.add_paragraph()
            op_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = op_para.add_run(ci.operator_name.upper())
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = DocColors.DARK_NAVY
            run.font.name = 'Calibri'

        self.doc.add_paragraph("")

        # Document Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("DRILLING PROGRAM")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = DocColors.NAVY
        run.font.name = 'Calibri'

        # Subtitle
        subtitle_para = self.doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run("& OPERATIONAL PROCEDURES")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = DocColors.STEEL_BLUE
        run.font.name = 'Calibri'

        self.doc.add_paragraph("")
        self._add_colored_line(DocColors.RED, width=2)
        self.doc.add_paragraph("")

        # Well Information Table on Cover
        cover_data = []
        if ci.well_name:
            cover_data.append(("Well Name", ci.well_name))
        if ci.well_number:
            cover_data.append(("Well Number", ci.well_number))
        if ci.field_name:
            cover_data.append(("Field", ci.field_name))
        if ci.country:
            cover_data.append(("Country / Region",
                               f"{ci.country} / {ci.region}"))
        if ci.rig_name:
            cover_data.append(("Rig", f"{ci.rig_name} ({ci.rig_type})"))
        if wi.well_type:
            cover_data.append(("Well Type", wi.well_type))
        if wi.well_profile:
            cover_data.append(("Well Profile", wi.well_profile))
        if wi.total_depth_md > 0:
            cover_data.append(("Planned TD (MD/TVD)",
                               f"{wi.total_depth_md:,.0f} ft / "
                               f"{wi.total_depth_tvd:,.0f} ft"))
        if wi.target_formation:
            cover_data.append(("Target Formation", wi.target_formation))

        if cover_data:
            table = self.doc.add_table(rows=len(cover_data), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            for i, (key, val) in enumerate(cover_data):
                row = table.rows[i]
                # Key cell
                cell_key = row.cells[0]
                cell_key.width = Inches(2.5)
                TableHelper.format_label_cell(cell_key, key, font_size=11)

                # Value cell
                cell_val = row.cells[1]
                cell_val.width = Inches(4.0)
                cell_val.text = ""
                p = cell_val.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(val))
                r.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = DocColors.DARK_NAVY
                r.font.name = 'Calibri'

        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        # Document Control Table
        control_data = [
            ("Document Number", ci.document_number or "DRL-PRG-001"),
            ("Revision", ci.revision or "0"),
            ("Date", datetime.now().strftime("%d-%B-%Y")),
            ("Prepared By", ci.prepared_by),
            ("Reviewed By", ci.reviewed_by),
            ("Approved By", ci.approved_by),
        ]

        ctrl_table = self.doc.add_table(rows=len(control_data), cols=2)
        ctrl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ctrl_table.style = 'Table Grid'

        for i, (key, val) in enumerate(control_data):
            row = ctrl_table.rows[i]
            cell_key = row.cells[0]
            cell_key.width = Inches(2.0)
            TableHelper.format_label_cell(cell_key, key, font_size=10)
            cell_val = row.cells[1]
            cell_val.width = Inches(4.5)
            TableHelper.format_value_cell(cell_val, val or "", font_size=10)

        # Contractor line
        self.doc.add_paragraph("")
        if ci.contractor_name:
            c_para = self.doc.add_paragraph()
            c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = c_para.add_run(
                f"Drilling Contractor: {ci.contractor_name}")
            run.font.size = Pt(10)
            run.font.color.rgb = DocColors.MEDIUM_GRAY
            run.font.name = 'Calibri'

        # Bottom border
        self._add_colored_line(DocColors.DARK_NAVY, width=3)

        # Disclaimer
        disc_para = self.doc.add_paragraph()
        disc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disc_para.add_run(
            f"This document is {ci.classification}. "
            f"Unauthorized distribution is prohibited."
        )
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = DocColors.RED

        self.doc.add_page_break()

    # ========================================================================
    # REVISION HISTORY
    # ========================================================================

    def _create_revision_history(self):
        """تاریخچه ویرایش"""
        self._add_section_heading("REVISION HISTORY", level=1)

        headers = ["Rev.", "Date", "Description of Changes",
                    "Prepared By", "Reviewed By", "Approved By"]
        data = [
            [self.project.company_info.revision or "0",
             datetime.now().strftime("%d-%b-%Y"),
             "Initial Issue - Issued for Review",
             self.project.company_info.prepared_by or "",
             self.project.company_info.reviewed_by or "",
             self.project.company_info.approved_by or ""],
        ]

        TableHelper.create_professional_table(
            self.doc, headers, data,
            col_widths=[0.5, 1.0, 2.5, 1.2, 1.2, 1.2]
        )

        self.doc.add_paragraph("")

        # Approval Signatures
        self._add_section_heading("APPROVAL SIGNATURES", level=2)

        sig_table = self.doc.add_table(rows=4, cols=4)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.style = 'Table Grid'

        sig_headers = ["Role", "Name", "Signature", "Date"]
        for i, h in enumerate(sig_headers):
            TableHelper.format_header_cell(
                sig_table.rows[0].cells[i], h)

        roles = [
            ("Prepared By (Drilling Engineer)",
             self.project.company_info.prepared_by),
            ("Reviewed By (Senior Drilling Engineer)",
             self.project.company_info.reviewed_by),
            ("Approved By (Drilling Manager)",
             self.project.company_info.approved_by),
        ]

        for i, (role, name) in enumerate(roles):
            row = sig_table.rows[i + 1]
            TableHelper.set_row_height(row, 1.5)
            TableHelper.format_data_cell(
                row.cells[0], role, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            TableHelper.format_data_cell(row.cells[1], name or "")
            TableHelper.format_data_cell(row.cells[2], "")
            TableHelper.format_data_cell(row.cells[3], "")

        self.doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS
    # ========================================================================

    def _create_table_of_contents(self):
        """فهرست مطالب"""
        self._add_section_heading("TABLE OF CONTENTS", level=1)

        toc_items = [
            ("1.", "EXECUTIVE SUMMARY", ""),
            ("2.", "WELL INFORMATION", ""),
            ("   2.1", "General Well Data", ""),
            ("   2.2", "Location & Coordinates", ""),
            ("   2.3", "Reservoir Information", ""),
            ("3.", "RIG SPECIFICATIONS", ""),
            ("4.", "FORMATION PROGNOSIS", ""),
            ("   4.1", "Formation Tops", ""),
            ("   4.2", "Pore Pressure & Fracture Gradient", ""),
            ("5.", "HAZARD ANALYSIS & RISK ASSESSMENT", ""),
            ("6.", "CASING DESIGN", ""),
            ("   6.1", "Casing Program Summary", ""),
            ("   6.2", "Design Basis & Calculations", ""),
            ("7.", "DRILLING FLUID PROGRAM", ""),
            ("   7.1", "Mud Program by Section", ""),
            ("   7.2", "Solids Control Equipment", ""),
            ("8.", "BHA & DRILL STRING DESIGN", ""),
            ("   8.1", "BHA Summary", ""),
            ("   8.2", "Drilling Parameters", ""),
            ("9.", "HYDRAULICS ANALYSIS", ""),
            ("10.", "CEMENTING PROGRAM", ""),
            ("   10.1", "Cement Design Summary", ""),
            ("   10.2", "Cement Volume Calculations", ""),
            ("11.", "DIRECTIONAL DRILLING PLAN", ""),
            ("   11.1", "Directional Summary", ""),
            ("   11.2", "Wellpath Data", ""),
            ("12.", "BOP & WELL CONTROL", ""),
            ("   12.1", "BOP Configuration", ""),
            ("   12.2", "Well Control Data", ""),
            ("13.", "EVALUATION PROGRAM", ""),
            ("14.", "TIME vs DEPTH ESTIMATE", ""),
            ("15.", "DETAILED OPERATING PROCEDURES", ""),
            ("   15.1", "Pre-Spud Checklist", ""),
            ("   15.2", "Drilling Procedures", ""),
            ("   15.3", "Casing Running Procedures", ""),
            ("   15.4", "Cementing Procedures", ""),
            ("   15.5", "Tripping Procedures", ""),
            ("   15.6", "BOP Procedures", ""),
            ("   15.7", "Well Control Procedures", ""),
            ("   15.8", "LOT/FIT Procedures", ""),
            ("   15.9", "Directional Drilling Procedures", ""),
            ("   15.10", "Logging Procedures", ""),
            ("   15.11", "Stuck Pipe Procedures", ""),
            ("   15.12", "Lost Circulation Procedures", ""),
            ("16.", "EMERGENCY PROCEDURES", ""),
            ("   16.1", "H₂S Emergency", ""),
            ("   16.2", "Fire Emergency", ""),
            ("   16.3", "Well Control Emergency", ""),
            ("17.", "HSE REQUIREMENTS", ""),
            ("18.", "APPENDICES", ""),
        ]

        table = self.doc.add_table(rows=len(toc_items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (num, title, page) in enumerate(toc_items):
            row = table.rows[i]
            # Number + Title
            cell = row.cells[0]
            cell.text = ""
            p = cell.paragraphs[0]
            indent = 0.5 if num.startswith("   ") else 0
            p.paragraph_format.left_indent = Cm(indent)

            run_num = p.add_run(num.strip() + "  ")
            run_num.bold = True if not num.startswith("   ") else False
            run_num.font.size = Pt(10) if not num.startswith("   ") else Pt(9)
            run_num.font.color.rgb = DocColors.DARK_NAVY
            run_num.font.name = 'Calibri'

            run_title = p.add_run(title)
            run_title.bold = True if not num.startswith("   ") else False
            run_title.font.size = Pt(10) if not num.startswith("   ") else Pt(9)
            run_title.font.color.rgb = DocColors.BLACK
            run_title.font.name = 'Calibri'

            # Separator line (dots)
            cell2 = row.cells[1]
            cell2.text = ""
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self.doc.add_page_break()

    # ========================================================================
    # ABBREVIATIONS
    # ========================================================================

    def _create_abbreviations(self):
        """اختصارات و تعاریف"""
        self._add_section_heading("ABBREVIATIONS & DEFINITIONS", level=1)

        abbreviations = [
            ("AFE", "Authorization for Expenditure"),
            ("API", "American Petroleum Institute"),
            ("BHA", "Bottom Hole Assembly"),
            ("BHP", "Bottom Hole Pressure"),
            ("BHT", "Bottom Hole Temperature"),
            ("BOP", "Blowout Preventer"),
            ("BPM", "Barrels Per Minute"),
            ("BUR", "Build Up Rate"),
            ("CBL", "Cement Bond Log"),
            ("CCL", "Casing Collar Locator"),
            ("CBIL", "Cement Bond Imaging Log"),
            ("DC", "Drill Collar"),
            ("DF", "Design Factor"),
            ("DLS", "Dog Leg Severity"),
            ("DP", "Drill Pipe"),
            ("DST", "Drill Stem Test"),
            ("ECD", "Equivalent Circulating Density"),
            ("EMW", "Equivalent Mud Weight"),
            ("ERD", "Extended Reach Drilling"),
            ("FIT", "Formation Integrity Test"),
            ("FL", "Fluid Loss"),
            ("FV", "Funnel Viscosity"),
            ("GPM", "Gallons Per Minute"),
            ("GR", "Gamma Ray"),
            ("HCR", "Hydraulic Control Remote"),
            ("HHP", "Hydraulic Horsepower"),
            ("HSE", "Health, Safety & Environment"),
            ("HSI", "Horsepower per Square Inch"),
            ("HTHP", "High Temperature High Pressure"),
            ("HWDP", "Heavy Weight Drill Pipe"),
            ("IADC", "International Association of Drilling Contractors"),
            ("ICP", "Initial Circulating Pressure"),
            ("KOP", "Kick-Off Point"),
            ("LCM", "Lost Circulation Material"),
            ("LOT", "Leak-Off Test"),
            ("LWD", "Logging While Drilling"),
            ("MAASP", "Maximum Allowable Annular Surface Pressure"),
            ("MD", "Measured Depth"),
            ("MW", "Mud Weight"),
            ("MWD", "Measurement While Drilling"),
            ("NACE", "National Association of Corrosion Engineers"),
            ("NPT", "Non-Productive Time"),
            ("OBM", "Oil-Based Mud"),
            ("OWR", "Oil-Water Ratio"),
            ("PDC", "Polycrystalline Diamond Compact"),
            ("POOH", "Pull Out Of Hole"),
            ("PV", "Plastic Viscosity"),
            ("PWD", "Pressure While Drilling"),
            ("RIH", "Run In Hole"),
            ("RKB", "Rotary Kelly Bushing"),
            ("ROP", "Rate of Penetration"),
            ("RPM", "Revolutions Per Minute"),
            ("RSS", "Rotary Steerable System"),
            ("SBM", "Synthetic-Based Mud"),
            ("SCBA", "Self-Contained Breathing Apparatus"),
            ("SICP", "Shut-In Casing Pressure"),
            ("SIDPP", "Shut-In Drill Pipe Pressure"),
            ("SPM", "Strokes Per Minute"),
            ("SPP", "Stand Pipe Pressure"),
            ("TD", "Total Depth"),
            ("TFA", "Total Flow Area"),
            ("TOC", "Top of Cement"),
            ("TOOH", "Trip Out Of Hole"),
            ("TVD", "True Vertical Depth"),
            ("VDL", "Variable Density Log"),
            ("VME", "Von Mises Equivalent"),
            ("WBM", "Water-Based Mud"),
            ("WOB", "Weight on Bit"),
            ("WOC", "Wait on Cement"),
            ("YP", "Yield Point"),
        ]

        # Create table with 4 columns (2 pairs)
        mid = len(abbreviations) // 2 + len(abbreviations) % 2
        left = abbreviations[:mid]
        right = abbreviations[mid:]

        table = self.doc.add_table(rows=mid + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Headers
        for i, h in enumerate(["Abbrev.", "Definition",
                                "Abbrev.", "Definition"]):
            TableHelper.format_header_cell(
                table.rows[0].cells[i], h, font_size=8)

        for i in range(mid):
            row = table.rows[i + 1]
            TableHelper.set_row_height(row, 0.5)

            if i < len(left):
                TableHelper.format_data_cell(
                    row.cells[0], left[i][0], font_size=8, bold=True)
                TableHelper.format_data_cell(
                    row.cells[1], left[i][1], font_size=8,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
            if i < len(right):
                TableHelper.format_data_cell(
                    row.cells[2], right[i][0], font_size=8, bold=True)
                TableHelper.format_data_cell(
                    row.cells[3], right[i][1], font_size=8,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)

            if i % 2 == 0:
                for c in range(4):
                    if i < len(left) or (c >= 2 and i < len(right)):
                        TableHelper.set_cell_shading(
                            row.cells[c], DocColors.ROW_ALT_BG)

        self.doc.add_page_break()

    # ========================================================================
    # EXECUTIVE SUMMARY
    # ========================================================================

    def _create_executive_summary(self):
        """خلاصه اجرایی"""
        ci = self.project.company_info
        wi = self.project.well_info
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. EXECUTIVE SUMMARY", level=1)

        # Summary paragraph
        summary = (
            f"This Drilling Program has been prepared for the drilling of "
            f"well {ci.well_name or '[Well Name]'} located in "
            f"{ci.field_name or '[Field]'} field, "
            f"{ci.country or '[Country]'}. "
            f"The well is a {wi.well_type} {wi.well_profile} well "
            f"targeting the {wi.target_formation or '[Formation]'} "
            f"formation at a planned total depth of "
            f"{wi.total_depth_md:,.0f} ft MD "
            f"({wi.total_depth_tvd:,.0f} ft TVD). "
        )

        p = self.doc.add_paragraph()
        run = p.add_run(summary)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

        # Well Purpose
        self._add_section_heading("Well Objectives", level=2)
        objectives = [
            f"Drill a {wi.well_type} well to {wi.target_formation or 'target'} formation.",
            f"Achieve planned TD of {wi.total_depth_md:,.0f} ft MD / "
            f"{wi.total_depth_tvd:,.0f} ft TVD.",
            "Obtain comprehensive formation evaluation data.",
            "Ensure safe operations with zero incidents.",
            "Minimize NPT and optimize drilling performance.",
        ]
        for obj in objectives:
            p = self.doc.add_paragraph(obj, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

        # Key Parameters Summary Table
        self._add_section_heading("Key Well Parameters", level=2)

        key_params = [
            ("Operator", ci.operator_name or ""),
            ("Well Name", ci.well_name or ""),
            ("Field", ci.field_name or ""),
            ("Well Type", wi.well_type),
            ("Well Profile", wi.well_profile),
            ("Total Depth (MD)", f"{wi.total_depth_md:,.0f} ft"),
            ("Total Depth (TVD)", f"{wi.total_depth_tvd:,.0f} ft"),
            ("Target Formation", wi.target_formation or ""),
            ("Rig", f"{ci.rig_name} ({ci.rig_type})"),
            ("KB Elevation", f"{wi.kb_elevation:,.1f} ft"),
        ]

        if wi.water_depth > 0:
            key_params.append(
                ("Water Depth", f"{wi.water_depth:,.0f} ft"))
        if wi.expected_h2s_concentration > 0:
            key_params.append(
                ("H₂S Expected", f"{wi.expected_h2s_concentration}%"))
        if wi.nace_required:
            key_params.append(("NACE Required", "Yes"))

        TableHelper.create_key_value_table(
            self.doc, key_params, num_cols=4)

        # Casing Program Summary
        if self.project.casing_design:
            self._add_section_heading("Casing Program Overview", level=2)

            headers = ["Section", "Hole Size", "Casing OD",
                       "Weight", "Grade", "Depth (MD)", "Depth (TVD)"]
            data = []
            for cd in self.project.casing_design:
                data.append([
                    cd.section_name,
                    f'{cd.hole_size}"',
                    f'{cd.casing_od}"',
                    f"{cd.casing_weight} ppf",
                    cd.casing_grade,
                    f"{cd.setting_depth_md:,.0f} ft",
                    f"{cd.setting_depth_tvd:,.0f} ft",
                ])

            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=f"Table {self._next_table()}: "
                        f"Casing Program Summary"
            )

        # Estimated Time Summary
        if self.project.time_estimates:
            self._add_section_heading("Estimated Time Summary", level=2)

            total_days = 0
            for te in self.project.time_estimates:
                total_days = max(total_days, te.cumulative_days)

            time_para = self.doc.add_paragraph()
            run = time_para.add_run(
                f"Estimated total well duration: "
                f"{total_days:.1f} days")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DocColors.DARK_NAVY

        self.doc.add_page_break()

    # ========================================================================
    # WELL INFORMATION SECTION
    # ========================================================================

    def _create_well_information_section(self):
        """بخش اطلاعات چاه"""
        self.section_counter += 1
        ci = self.project.company_info
        wi = self.project.well_info

        self._add_section_heading(
            f"{self.section_counter}. WELL INFORMATION", level=1)

        # 2.1 General Well Data
        self._add_section_heading(
            f"{self.section_counter}.1 General Well Data", level=2)

        general_data = [
            ("Operator", ci.operator_name),
            ("Drilling Contractor", ci.contractor_name),
            ("Well Name", ci.well_name),
            ("Well Number", ci.well_number),
            ("Field Name", ci.field_name),
            ("Pad / Platform", ci.pad_name),
            ("Block / License", ci.block_license),
            ("API / Well ID", ci.api_number),
            ("Country", ci.country),
            ("Region / Area", ci.region),
            ("Rig Name", ci.rig_name),
            ("Rig Type", ci.rig_type),
            ("Well Type", wi.well_type),
            ("Well Profile", wi.well_profile),
            ("Planned Spud Date", ci.spud_date),
            ("Total Depth (MD)", f"{wi.total_depth_md:,.0f} ft"),
            ("Total Depth (TVD)", f"{wi.total_depth_tvd:,.0f} ft"),
            ("KB Elevation", f"{wi.kb_elevation:,.1f} ft"),
            ("Ground Elevation",
             f"{wi.ground_elevation:,.1f} ft"),
            ("Air Gap (RKB to MSL)",
             f"{wi.air_gap:,.1f} ft"),
        ]

        if wi.water_depth > 0:
            general_data.append(
                ("Water Depth", f"{wi.water_depth:,.0f} ft"))

        TableHelper.create_key_value_table(
            self.doc, general_data, num_cols=4)

        self.doc.add_paragraph("")

        # 2.2 Location & Coordinates
        self._add_section_heading(
            f"{self.section_counter}.2 Location & Coordinates", level=2)

        location_data = [
            ("Surface Latitude", wi.surface_latitude),
            ("Surface Longitude", wi.surface_longitude),
            ("Target Latitude", wi.target_latitude),
            ("Target Longitude", wi.target_longitude),
            ("Coordinate System", wi.coordinate_system),
            ("Magnetic Declination",
             f"{wi.magnetic_declination:+.2f}°"),
            ("Grid Convergence",
             f"{wi.grid_convergence:+.2f}°"),
        ]

        TableHelper.create_key_value_table(
            self.doc, location_data, num_cols=4)

        self.doc.add_paragraph("")

        # 2.3 Reservoir Information
        self._add_section_heading(
            f"{self.section_counter}.3 Target / Reservoir Information",
            level=2)

        reservoir_data = [
            ("Target Formation", wi.target_formation),
            ("Target Zone", wi.target_zone),
            ("Expected Reservoir Pressure",
             f"{wi.expected_reservoir_pressure:,.0f} psi"),
            ("Expected Reservoir Temperature",
             f"{wi.expected_reservoir_temperature:,.0f} °F"),
            ("Expected H₂S",
             f"{wi.expected_h2s_concentration:.1f} %"),
            ("Expected CO₂",
             f"{wi.expected_co2_concentration:.1f} %"),
            ("NACE MR-0175 Required",
             "Yes" if wi.nace_required else "No"),
            ("Wellhead Type", wi.wellhead_type),
            ("Christmas Tree Type", wi.xmas_tree_type),
        ]

        TableHelper.create_key_value_table(
            self.doc, reservoir_data, num_cols=4)

        if wi.seismic_reference:
            self.doc.add_paragraph("")
            self._add_section_heading("Seismic Reference", level=3)
            p = self.doc.add_paragraph(wi.seismic_reference)
            for run in p.runs:
                run.font.size = Pt(10)

        self.doc.add_page_break()

    # ========================================================================
    # RIG SPECIFICATIONS
    # ========================================================================

    def _create_rig_specification_section(self):
        """بخش مشخصات دکل"""
        self.section_counter += 1
        rs = self.project.rig_spec

        self._add_section_heading(
            f"{self.section_counter}. RIG SPECIFICATIONS", level=1)

        rig_data = [
            ("Rig Name", rs.rig_name),
            ("Rig Type", rs.rig_type),
            ("Rig Contractor", rs.rig_contractor),
            ("Max Hook Load",
             f"{rs.max_hook_load:,.0f} lbs"),
            ("Drawworks Power",
             f"{rs.drawworks_power:,.0f} HP"),
            ("Top Drive",
             f"{'Yes' if rs.top_drive else 'No'} - {rs.top_drive_model}"),
            ("Top Drive Torque",
             f"{rs.top_drive_torque:,.0f} ft-lbs"),
            ("Max Rotary Speed",
             f"{rs.max_rotary_speed:,.0f} RPM"),
            ("Derrick Height",
             f"{rs.derrick_height:,.0f} ft"),
            ("Rotary Table",
             f'{rs.rotary_table_size}" opening'),
        ]

        TableHelper.create_key_value_table(
            self.doc, rig_data, num_cols=4)

        # Mud Pumps
        self._add_section_heading("Mud Pumps", level=2)

        pump_headers = ["Parameter", "Pump #1", "Pump #2", "Pump #3"]
        pump_data = [
            ["Type", rs.mud_pump_1_type,
             rs.mud_pump_2_type, rs.mud_pump_3_type],
            ["HP", f"{rs.mud_pump_1_hp:,.0f}",
             f"{rs.mud_pump_2_hp:,.0f}",
             f"{rs.mud_pump_3_hp:,.0f}"],
            ["Liner Size",
             f'{rs.mud_pump_1_liner}"',
             f'{rs.mud_pump_2_liner}"', ""],
            ["Max Pressure",
             f"{rs.mud_pump_1_max_pressure:,.0f} psi",
             f"{rs.mud_pump_2_max_pressure:,.0f} psi", ""],
            ["Max Flow",
             f"{rs.mud_pump_1_max_flow:,.0f} GPM",
             f"{rs.mud_pump_2_max_flow:,.0f} GPM", ""],
        ]

        TableHelper.create_professional_table(
            self.doc, pump_headers, pump_data,
            caption=f"Table {self._next_table()}: Mud Pump Specifications"
        )

        # Pit System
        self._add_section_heading("Mud System & Utilities", level=2)

        pit_data = [
            ("Total Pit Volume",
             f"{rs.pit_volume_total:,.0f} bbl"),
            ("Active Pit Volume",
             f"{rs.pit_volume_active:,.0f} bbl"),
            ("Shale Shakers",
             f"{rs.shale_shaker_count} units"),
            ("Degasser", rs.degasser_type),
            ("Centrifuge", rs.centrifuge),
            ("Generators", rs.generators),
            ("Total Power",
             f"{rs.total_power:,.0f} kW"),
            ("Crane Capacity",
             f"{rs.crane_capacity:,.0f} tons"),
            ("Accommodation",
             f"{rs.accommodation} persons"),
        ]

        TableHelper.create_key_value_table(
            self.doc, pit_data, num_cols=4)

        self.doc.add_page_break()

    # ========================================================================
    # FORMATION PROGNOSIS
    # ========================================================================

    def _create_formation_prognosis_section(self):
        """بخش پیش‌بینی سازندها"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. FORMATION PROGNOSIS", level=1)

        if not self.project.formation_tops:
            self.doc.add_paragraph(
                "No formation data entered. "
                "Refer to geological prognosis report."
            )
            self.doc.add_page_break()
            return

        # Formation Tops Table
        self._add_section_heading("Formation Tops", level=2)

        headers = [
            "Formation", "Lithology",
            "MD Top\n(ft)", "MD Bot\n(ft)",
            "TVD Top\n(ft)", "TVD Bot\n(ft)",
            "PP Top\n(ppg)", "PP Bot\n(ppg)",
            "FG Top\n(ppg)", "FG Bot\n(ppg)",
            "Temp\n(°F)", "Remarks"
        ]

        data = []
        for ft in self.project.formation_tops:
            data.append([
                ft.name,
                ft.formation_type,
                f"{ft.md_top:,.0f}",
                f"{ft.md_bottom:,.0f}",
                f"{ft.tvd_top:,.0f}",
                f"{ft.tvd_bottom:,.0f}",
                f"{ft.pore_pressure_top:.1f}",
                f"{ft.pore_pressure_bottom:.1f}",
                f"{ft.fracture_gradient_top:.1f}",
                f"{ft.fracture_gradient_bottom:.1f}",
                f"{ft.temperature_bottom:.0f}",
                ft.remarks or "",
            ])

        TableHelper.create_professional_table(
            self.doc, headers, data,
            caption=f"Table {self._next_table()}: Formation Tops Prognosis"
        )

        # Formation Description
        self._add_section_heading("Formation Descriptions", level=2)

        for ft in self.project.formation_tops:
            p = self.doc.add_paragraph()
            run = p.add_run(f"{ft.name} ({ft.formation_type}): ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DocColors.DARK_NAVY

            desc = (
                f"Expected from {ft.md_top:,.0f} ft to "
                f"{ft.md_bottom:,.0f} ft MD. "
                f"Pore pressure: {ft.pore_pressure_top:.1f} - "
                f"{ft.pore_pressure_bottom:.1f} ppg EMW. "
                f"Fracture gradient: {ft.fracture_gradient_top:.1f} - "
                f"{ft.fracture_gradient_bottom:.1f} ppg EMW. "
                f"Temperature: {ft.temperature_top:.0f} - "
                f"{ft.temperature_bottom:.0f} °F. "
                f"Drillability: {ft.drillability}. "
            )
            if ft.remarks:
                desc += f"Note: {ft.remarks}."

            run2 = p.add_run(desc)
            run2.font.size = Pt(10)

        self.doc.add_page_break()

    # ========================================================================
    # HAZARD ANALYSIS
    # ========================================================================

    def _create_hazard_analysis_section(self):
        """بخش تحلیل خطرات"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. HAZARD ANALYSIS & RISK ASSESSMENT",
            level=1)

        if not self.project.hazards:
            # Generate default hazards based on formations
            self._add_default_hazard_discussion()
        else:
            headers = [
                "Hazard", "Depth\n(ft MD)",
                "Severity", "Probability",
                "Description",
                "Mitigation", "Contingency"
            ]

            data = []
            for h in self.project.hazards:
                data.append([
                    h.hazard_type,
                    f"{h.md_top:,.0f}-{h.md_bottom:,.0f}",
                    h.severity,
                    h.probability,
                    h.description,
                    h.mitigation,
                    h.contingency,
                ])

            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=f"Table {self._next_table()}: "
                        f"Hazard Analysis Summary"
            )

        # Risk Matrix
        self._add_section_heading("Risk Assessment Matrix", level=2)
        self._create_risk_matrix()

        self.doc.add_page_break()

    def _add_default_hazard_discussion(self):
        """بحث پیش‌فرض خطرات"""
        default_hazards = [
            ("Shallow Gas", "Risk of shallow gas pockets in "
             "unconsolidated surface formations. "
             "Mitigation: Diverter system installed, "
             "controlled ROP in surface section."),
            ("Lost Circulation", "Possible losses in fractured "
             "or vugular formations. "
             "Mitigation: LCM on location, cement on standby, "
             "monitor ECD."),
            ("Stuck Pipe", "Risk of differential sticking in "
             "permeable zones, mechanical sticking in deviated holes. "
             "Mitigation: Proper mud properties, minimize static time, "
             "jar in BHA."),
            ("Well Control", "Kick risk when penetrating pressured zones. "
             "Mitigation: Proper mud weight, BOP tested, "
             "well control drills conducted."),
            ("H₂S Gas", "If H₂S is expected, all NACE compliant "
             "materials required. "
             "Mitigation: H₂S detection, SCBA available, "
             "safety briefings."),
        ]

        for title, desc in default_hazards:
            p = self.doc.add_paragraph()
            run = p.add_run(f"• {title}: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DocColors.DARK_NAVY

            run2 = p.add_run(desc)
            run2.font.size = Pt(10)

    def _create_risk_matrix(self):
        """ماتریس ریسک"""
        table = self.doc.add_table(rows=6, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        headers = ["Severity →\nProbability ↓",
                    "Negligible", "Minor", "Moderate",
                    "Major", "Catastrophic"]
        for i, h in enumerate(headers):
            TableHelper.format_header_cell(
                table.rows[0].cells[i], h, font_size=8)

        probs = ["Almost Certain", "Likely",
                 "Possible", "Unlikely", "Rare"]

        # Risk colors (simplified)
        risk_colors = {
            'L': "27AE60",  # Green - Low
            'M': "F1C40F",  # Yellow - Medium
            'H': "E67E22",  # Orange - High
            'E': "E74C3C",  # Red - Extreme
        }

        risk_grid = [
            ['M', 'H', 'H', 'E', 'E'],  # Almost Certain
            ['M', 'M', 'H', 'H', 'E'],  # Likely
            ['L', 'M', 'M', 'H', 'E'],  # Possible
            ['L', 'L', 'M', 'M', 'H'],  # Unlikely
            ['L', 'L', 'L', 'M', 'M'],  # Rare
        ]

        risk_labels = {
            'L': 'LOW', 'M': 'MEDIUM',
            'H': 'HIGH', 'E': 'EXTREME'
        }

        for row_idx, prob in enumerate(probs):
            row = table.rows[row_idx + 1]
            TableHelper.set_row_height(row, 0.7)

            # Probability label
            TableHelper.format_subheader_cell(
                row.cells[0], prob, font_size=8)

            for col_idx in range(5):
                cell = row.cells[col_idx + 1]
                risk = risk_grid[row_idx][col_idx]
                color = risk_colors[risk]
                TableHelper.set_cell_shading(cell, color)
                TableHelper.format_data_cell(
                    cell, risk_labels[risk], font_size=8, bold=True,
                    color=DocColors.WHITE if risk in ['H', 'E']
                    else DocColors.BLACK
                )

        # Legend
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for risk, label in [('L', 'Low'), ('M', 'Medium'),
                            ('H', 'High'), ('E', 'Extreme')]:
            run = p.add_run(f"  ■ {label}  ")
            run.font.size = Pt(9)
            run.font.bold = True

    # ========================================================================
    # CASING DESIGN SECTION
    # ========================================================================

    def _create_casing_design_section(self):
        """بخش طراحی لوله جداری"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. CASING DESIGN", level=1)

        if not self.project.casing_design:
            self.doc.add_paragraph("No casing design data entered.")
            self.doc.add_page_break()
            return

        # Casing Program Summary
        self._add_section_heading("Casing Program Summary", level=2)

        headers = [
            "Section", "Hole Size\n(in)", "Casing OD\n(in)",
            "Casing ID\n(in)", "Weight\n(ppf)", "Grade",
            "Connection", "Setting Depth\nMD (ft)",
            "Setting Depth\nTVD (ft)", "TOC\nMD (ft)"
        ]

        data = []
        for cd in self.project.casing_design:
            data.append([
                cd.section_name,
                f"{cd.hole_size}",
                f"{cd.casing_od}",
                f"{cd.casing_id}",
                f"{cd.casing_weight}",
                cd.casing_grade,
                cd.casing_connection,
                f"{cd.setting_depth_md:,.0f}",
                f"{cd.setting_depth_tvd:,.0f}",
                f"{cd.top_of_cement_md:,.0f}",
            ])

        TableHelper.create_professional_table(
            self.doc, headers, data,
            caption=f"Table {self._next_table()}: "
                    f"Casing Program Summary"
        )

        # Design Factors
        self._add_section_heading("Design Factors", level=2)

        df_headers = [
            "Section", "Casing\n(OD x Wt x Grade)",
            "Burst\nRating\n(psi)", "Collapse\nRating\n(psi)",
            "Tensile\nRating\n(klbs)",
            "Min DF\nBurst", "Min DF\nCollapse", "Min DF\nTension"
        ]

        df_data = []
        for cd in self.project.casing_design:
            df_data.append([
                cd.section_name,
                f'{cd.casing_od}" x {cd.casing_weight} x {cd.casing_grade}',
                f"{cd.burst_rating:,.0f}",
                f"{cd.collapse_rating:,.0f}",
                f"{cd.tensile_rating:,.0f}",
                f"{cd.min_design_factor_burst:.2f}",
                f"{cd.min_design_factor_collapse:.2f}",
                f"{cd.min_design_factor_tension:.2f}",
            ])

        TableHelper.create_professional_table(
            self.doc, df_headers, df_data,
            caption=f"Table {self._next_table()}: "
                    f"Casing Design Factors"
        )

        # Casing Accessories
        self._add_section_heading("Casing Accessories", level=2)

        acc_headers = [
            "Section", "Float Shoe", "Float Collar\nDepth (ft)",
            "Centralizer\nType", "Centralizer\nSpacing (ft)",
            "Liner", "Remarks"
        ]

        acc_data = []
        for cd in self.project.casing_design:
            acc_data.append([
                cd.section_name,
                cd.float_shoe_type,
                f"{cd.float_collar_depth:,.0f}",
                cd.centralizer_type,
                f"{cd.centralizer_spacing:.0f}",
                "Yes" if cd.is_liner else "No",
                cd.remarks or "",
            ])

        TableHelper.create_professional_table(
            self.doc, acc_headers, acc_data,
            caption=f"Table {self._next_table()}: "
                    f"Casing Accessories"
        )

        # Design Basis Notes
        self._add_section_heading("Design Basis", level=2)

        design_notes = [
            "Casing design performed per API 5C3 / ISO 10400 standards.",
            f"Minimum burst design factor: "
            f"{self.project.casing_design[0].min_design_factor_burst:.2f}",
            f"Minimum collapse design factor: "
            f"{self.project.casing_design[0].min_design_factor_collapse:.2f}",
            f"Minimum tension design factor: "
            f"{self.project.casing_design[0].min_design_factor_tension:.2f}",
            "Burst load case: Gas kick at shoe / full evacuation.",
            "Collapse load case: Full/partial evacuation.",
            "Tension load case: Running in mud + overpull + bending.",
            "Biaxial (VME) check applied where applicable.",
        ]

        for note in design_notes:
            p = self.doc.add_paragraph(note, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        self.doc.add_page_break()

    # ========================================================================
    # MUD PROGRAM SECTION
    # ========================================================================

    def _create_mud_program_section(self):
        """بخش برنامه سیال حفاری"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. DRILLING FLUID PROGRAM", level=1)

        if not self.project.mud_programs:
            self.doc.add_paragraph("No mud program data entered.")
            self.doc.add_page_break()
            return

        # Mud Program Summary
        self._add_section_heading("Mud Program Summary", level=2)

        headers = [
            "Section", "Hole\n(in)", "Depth\n(ft)",
            "Mud Type", "MW In\n(ppg)", "MW Out\n(ppg)",
            "FV\n(s/qt)", "PV\n(cP)", "YP\n(lb/100ft²)",
            "FL\n(ml)", "Vol\n(bbl)"
        ]

        data = []
        for mp in self.project.mud_programs:
            data.append([
                mp.section_name,
                f"{mp.hole_size}",
                f"{mp.depth_from:,.0f}-{mp.depth_to:,.0f}",
                mp.mud_type,
                f"{mp.mud_weight_in:.1f}",
                f"{mp.mud_weight_out:.1f}",
                f"{mp.funnel_viscosity:.0f}",
                f"{mp.plastic_viscosity:.0f}",
                f"{mp.yield_point:.0f}",
                f"{mp.fluid_loss:.1f}",
                f"{mp.total_volume_required:,.0f}",
            ])

        TableHelper.create_professional_table(
            self.doc, headers, data,
            caption=f"Table {self._next_table()}: "
                    f"Drilling Fluid Program Summary"
        )

        # Detailed Properties
        self._add_section_heading(
            "Detailed Mud Properties by Section", level=2)

        for mp in self.project.mud_programs:
            self._add_section_heading(
                f"{mp.section_name} Section - {mp.mud_type}",
                level=3)

            props = [
                ("Section", mp.section_name),
                ("Hole Size", f'{mp.hole_size}"'),
                ("Depth Range",
                 f"{mp.depth_from:,.0f} - {mp.depth_to:,.0f} ft"),
                ("Mud Type", mp.mud_type),
                ("Mud Weight In", f"{mp.mud_weight_in:.1f} ppg"),
                ("Mud Weight Out", f"{mp.mud_weight_out:.1f} ppg"),
                ("Funnel Viscosity",
                 f"{mp.funnel_viscosity:.0f} sec/qt"),
                ("Plastic Viscosity",
                 f"{mp.plastic_viscosity:.0f} cP"),
                ("Yield Point",
                 f"{mp.yield_point:.0f} lb/100ft²"),
                ("Gel Strength (10s/10m/30m)",
                 f"{mp.gel_strength_10s:.0f} / "
                 f"{mp.gel_strength_10m:.0f} / "
                 f"{mp.gel_strength_30m:.0f}"),
                ("API Fluid Loss",
                 f"{mp.fluid_loss:.1f} ml/30min"),
                ("HTHP Fluid Loss",
                 f"{mp.hthp_fluid_loss:.1f} ml"),
                ("pH", f"{mp.ph:.1f}"),
                ("Chlorides", f"{mp.chlorides:,.0f} ppm"),
                ("MBT", f"{mp.mbt:.0f} lb/bbl"),
                ("Sand Content", f"{mp.sand_content:.1f} %"),
                ("Total Volume Required",
                 f"{mp.total_volume_required:,.0f} bbl"),
                ("Active Volume",
                 f"{mp.active_volume:,.0f} bbl"),
                ("ECD at Shoe",
                 f"{mp.ecd_at_shoe:.2f} ppg"),
                ("ECD at TD", f"{mp.ecd_at_td:.2f} ppg"),
            ]

            if mp.oil_water_ratio:
                props.append(("Oil/Water Ratio", mp.oil_water_ratio))
            if mp.electrical_stability > 0:
                props.append(("Electrical Stability",
                              f"{mp.electrical_stability:.0f} volts"))
            if mp.key_additives:
                props.append(("Key Additives", mp.key_additives))
            if mp.remarks:
                props.append(("Remarks", mp.remarks))

            TableHelper.create_key_value_table(
                self.doc, props, num_cols=4)

            self.doc.add_paragraph("")

        self.doc.add_page_break()

    # ========================================================================
    # BHA SECTION
    # ========================================================================

    def _create_bha_section(self):
        """بخش BHA و رشته حفاری"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. BHA & DRILL STRING DESIGN",
            level=1)

        if not self.project.bha_designs:
            self.doc.add_paragraph("No BHA design data entered.")
            self.doc.add_page_break()
            return

        # BHA Summary
        self._add_section_heading("BHA Summary", level=2)

        for bha in self.project.bha_designs:
            self._add_section_heading(
                f"BHA #{bha.bha_number} - {bha.section_name} Section "
                f"({bha.hole_size}\" Hole)",
                level=3)

            bha_info = [
                ("Section", bha.section_name),
                ("BHA Number", str(bha.bha_number)),
                ("Hole Size", f'{bha.hole_size}"'),
                ("BHA Type", bha.bha_type),
                ("Bit Type", bha.bit_type),
                ("Bit Size", f'{bha.bit_size}"'),
                ("Bit Manufacturer", bha.bit_manufacturer),
                ("Bit Model", bha.bit_model),
                ("Nozzles (TFA)", bha.bit_nozzles),
            ]

            if bha.motor_type and bha.motor_type != '-':
                bha_info.extend([
                    ("Motor Type", bha.motor_type),
                    ("Motor OD", f'{bha.motor_od}"'),
                    ("Motor Bend", f"{bha.motor_bend}°"),
                ])

            if bha.rss_type and bha.rss_type != '-':
                bha_info.append(("RSS Type", bha.rss_type))

            if bha.mwd_type:
                bha_info.append(("MWD", bha.mwd_type))
            if bha.lwd_sensors:
                bha_info.append(("LWD Sensors", bha.lwd_sensors))

            bha_info.extend([
                ("Recommended WOB", bha.recommended_wob),
                ("Recommended RPM", bha.recommended_rpm),
                ("Recommended Flow Rate",
                 bha.recommended_flow_rate),
                ("Max Torque", bha.recommended_torque),
            ])

            if bha.remarks:
                bha_info.append(("Remarks", bha.remarks))

            TableHelper.create_key_value_table(
                self.doc, bha_info, num_cols=4)

            self.doc.add_paragraph("")

        # Drilling Parameters
        if self.project.drilling_parameters:
            self._add_section_heading(
                "Drilling Parameters by Section", level=2)

            headers = [
                "Section", "Hole\n(in)",
                "WOB\n(klbs)", "RPM",
                "Flow\n(GPM)", "Max Torque\n(ft-lbs)",
                "ROP Avg\n(ft/hr)", "Max ECD\n(ppg)"
            ]

            data = []
            for dp in self.project.drilling_parameters:
                data.append([
                    dp.section_name,
                    f"{dp.hole_size}",
                    f"{dp.wob_min}-{dp.wob_max}",
                    f"{dp.rpm_min:.0f}-{dp.rpm_max:.0f}",
                    f"{dp.flow_rate_min:.0f}-{dp.flow_rate_max:.0f}",
                    f"{dp.torque_max:,.0f}",
                    f"{dp.rop_average:.0f}",
                    f"{dp.max_ecd:.1f}" if dp.max_ecd > 0 else "Monitor",
                ])

            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=f"Table {self._next_table()}: "
                        f"Drilling Parameters Summary"
            )

        self.doc.add_page_break()

    # ========================================================================
    # HYDRAULICS SECTION
    # ========================================================================

    def _create_hydraulics_section(self):
        """بخش هیدرولیک"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. HYDRAULICS ANALYSIS", level=1)

        p = self.doc.add_paragraph()
        run = p.add_run(
            "Hydraulics calculations are performed per API RP 13D. "
            "The following table summarizes the hydraulic parameters "
            "for each hole section. "
            "Detailed hydraulic analysis including ECD management, "
            "hole cleaning parameters, and surge/swab analysis "
            "should be performed by the drilling engineer prior to "
            "each section."
        )
        run.font.size = Pt(10)

        if self.project.mud_programs:
            headers = [
                "Section", "Hole Size\n(in)",
                "MW\n(ppg)", "ECD Shoe\n(ppg)",
                "ECD TD\n(ppg)"
            ]

            data = []
            for mp in self.project.mud_programs:
                data.append([
                    mp.section_name,
                    f"{mp.hole_size}",
                    f"{mp.mud_weight_out:.1f}",
                    f"{mp.ecd_at_shoe:.2f}" if mp.ecd_at_shoe > 0 else "TBD",
                    f"{mp.ecd_at_td:.2f}" if mp.ecd_at_td > 0 else "TBD",
                ])

            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=f"Table {self._next_table()}: "
                        f"Hydraulics Summary"
            )

        # Hydraulic Guidelines
        self._add_section_heading("Hydraulic Guidelines", level=2)

        guidelines = [
            "Maintain minimum annular velocity of 120 ft/min "
            "in vertical sections.",
            "Increase annular velocity to 150-180 ft/min in "
            "deviated/horizontal sections.",
            "ECD shall not exceed fracture gradient at any point.",
            "Pump hi-vis sweeps regularly in deviated sections.",
            "Monitor surge/swab during all tripping operations.",
            "Maximum trip speed shall be calculated to avoid "
            "surge/swab issues.",
            "HSI at bit should be maintained above 2.0 hp/sq.in.",
            "Bit nozzle selection should optimize hole cleaning.",
        ]

        for g in guidelines:
            p = self.doc.add_paragraph(g, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        self.doc.add_page_break()

    # ========================================================================
    # CEMENT SECTION
    # ========================================================================

    def _create_cement_section(self):
        """بخش سیمانکاری"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. CEMENTING PROGRAM", level=1)

        if not self.project.cement_design:
            self.doc.add_paragraph("No cement design data entered.")
            self.doc.add_page_break()
            return

        # Cement Program Summary
        self._add_section_heading("Cement Program Summary", level=2)

        for cd in self.project.cement_design:
            self._add_section_heading(
                f"{cd.section_name} Section Cementing", level=3)

            cement_data = [
                ("Section", cd.section_name),
                ("Casing OD", f'{cd.casing_od}"'),
                ("Hole Size", f'{cd.hole_size}"'),
                ("Shoe Depth (MD)", f"{cd.shoe_depth_md:,.0f} ft"),
                ("TOC (MD)", f"{cd.toc_md:,.0f} ft"),
                ("Lead Slurry Type", cd.lead_slurry_type),
                ("Lead Density", f"{cd.lead_slurry_density:.1f} ppg"),
                ("Lead Volume", f"{cd.lead_slurry_volume:.0f} bbl"),
                ("Lead Thickening Time",
                 f"{cd.lead_slurry_thickening_time:.0f} hrs"),
                ("Lead Comp. Strength (24hr)",
                 f"{cd.lead_slurry_compressive_strength:,.0f} psi"),
                ("Tail Slurry Type", cd.tail_slurry_type),
                ("Tail Density", f"{cd.tail_slurry_density:.1f} ppg"),
                ("Tail Volume", f"{cd.tail_slurry_volume:.0f} bbl"),
                ("Tail Thickening Time",
                 f"{cd.tail_slurry_thickening_time:.0f} hrs"),
                ("Tail Comp. Strength (24hr)",
                 f"{cd.tail_slurry_compressive_strength:,.0f} psi"),
                ("Spacer Type", cd.spacer_type),
                ("Spacer Volume", f"{cd.spacer_volume:.0f} bbl"),
                ("Displacement Volume",
                 f"{cd.displacement_volume:.0f} bbl"),
                ("Displacement Rate",
                 f"{cd.displacement_rate:.0f} bpm"),
                ("Excess", f"{cd.excess_percentage:.0f} %"),
                ("WOC Time", f"{cd.woc_time:.0f} hrs"),
                ("Plug Bump Pressure",
                 f"{cd.plug_bump_pressure:,.0f} psi"),
                ("CBL/CBIL Required",
                 "Yes" if cd.cbl_cbil_required else "No"),
            ]

            if cd.cement_additives:
                cement_data.append(
                    ("Cement Additives", cd.cement_additives))
            if cd.remarks:
                cement_data.append(("Remarks", cd.remarks))

            TableHelper.create_key_value_table(
                self.doc, cement_data, num_cols=4)

            self.doc.add_paragraph("")

        # Cementing Best Practices
        self._add_section_heading("Cementing Best Practices", level=2)

        practices = [
            "Condition hole and mud before cementing - "
            "circulate minimum 2 x bottoms up.",
            "Reduce mud properties (YP, gels) to minimum "
            "for cement displacement.",
            "Maintain minimum contact time of 10 minutes "
            "across critical intervals.",
            "Ensure spacer density is between mud and "
            "cement density.",
            "Reciprocate or rotate casing during cement "
            "placement where possible.",
            "Monitor ECD during cementing - do not exceed "
            "fracture gradient.",
            "Verify cement density at manifold throughout job.",
            "Bump plug and hold pressure to verify seal.",
            "Perform CBL/CBIL logging after WOC to verify "
            "cement quality.",
        ]

        for practice in practices:
            p = self.doc.add_paragraph(practice, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        self.doc.add_page_break()

    # ========================================================================
    # HELPER METHODS
    # ========================================================================

    def _add_section_heading(self, text: str, level: int = 1):
        """اضافه کردن هدر سکشن با فرمت حرفه‌ای"""
        if level == 1:
            # Full width colored heading
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

            # Add background shading to paragraph
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{DocColors.HEADER_BG}" '
                f'w:val="clear"/>'
            )
            pPr.append(shading)

            run = p.add_run(f"  {text}")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = DocColors.WHITE
            run.font.name = 'Calibri'

        elif level == 2:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

            # Border bottom
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" '
                f'w:space="1" w:color="1B4F72"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = DocColors.STEEL_BLUE
            run.font.name = 'Calibri'

        elif level == 3:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DocColors.MEDIUM_BLUE
            run.font.name = 'Calibri'

    def _add_colored_line(self, color: RGBColor, width: int = 2):
        """اضافه کردن خط رنگی"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        pPr = p._p.get_or_add_pPr()
        color_hex = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="{width * 8}" '
            f'w:space="1" w:color="{color_hex}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _next_table(self) -> int:
        """شماره جدول بعدی"""
        self.table_counter += 1
        return self.table_counter

    def _next_figure(self) -> int:
        """شماره شکل بعدی"""
        self.figure_counter += 1
        return self.figure_counter

    def _add_note(self, text: str, note_type: str = "NOTE"):
        """اضافه کردن یادداشت"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)

        # Icon and label
        if note_type == "WARNING":
            label = "⚠️ WARNING: "
            color = DocColors.RED
            bg = DocColors.CAUTION_BG
        elif note_type == "CAUTION":
            label = "⚠ CAUTION: "
            color = DocColors.ORANGE
            bg = DocColors.WARNING_BG
        elif note_type == "IMPORTANT":
            label = "❗ IMPORTANT: "
            color = DocColors.DARK_RED
            bg = DocColors.CAUTION_BG
        else:
            label = "📌 NOTE: "
            color = DocColors.STEEL_BLUE
            bg = DocColors.INFO_BG

        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = color
        run_label.font.name = 'Calibri'

        run_text = p.add_run(text)
        run_text.font.size = Pt(9)
        run_text.font.color.rgb = DocColors.BLACK
        run_text.font.name = 'Calibri'

    def _add_procedure_steps(self, steps: List[str]):
        """اضافه کردن مراحل پروسیجر"""
        for step in steps:
            if not step.strip():
                self.doc.add_paragraph("")
                continue

            # Determine indent level and formatting
            stripped = step.lstrip()
            indent_level = len(step) - len(stripped)

            if stripped.startswith(("1.", "2.", "3.", "4.", "5.",
                                     "6.", "7.", "8.", "9.", "10.",
                                     "11.", "12.", "13.", "14.", "15.",
                                     "16.", "17.", "18.", "19.", "20.")):
                # Main section header in procedure
                if indent_level == 0 and not any(
                    c.isdigit() and step[step.index(c)+1:step.index(c)+2] == '.'
                    for c in step[:3] if c.isdigit()
                ):
                    p = self.doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = DocColors.DARK_NAVY
                    run.font.name = 'Calibri'
                else:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.3)
                    p.paragraph_format.space_after = Pt(2)

                    # Split into number and text
                    parts = stripped.split(' ', 1)
                    if len(parts) == 2:
                        run_num = p.add_run(parts[0] + " ")
                        run_num.bold = True
                        run_num.font.size = Pt(10)
                        run_num.font.color.rgb = DocColors.NAVY
                        run_num.font.name = 'Calibri'

                        run_text = p.add_run(parts[1])
                        run_text.font.size = Pt(10)
                        run_text.font.name = 'Calibri'
                    else:
                        run = p.add_run(stripped)
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'

            elif stripped.upper() == stripped and len(stripped) > 5 \
                    and not stripped.startswith(("   ", "-", "•", "a)", "b)")):
                # Section title in procedure
                self._add_section_heading(stripped, level=3)

            elif stripped.startswith(("a)", "b)", "c)", "d)", "e)",
                                       "f)", "g)", "h)", "i)",
                                       "- ", "• ", "* ")):
                # Sub-bullet
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5 + indent_level * 0.2)
                p.paragraph_format.space_after = Pt(1)

                run = p.add_run(stripped)
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                run.font.color.rgb = DocColors.DARK_GRAY

            else:
                # Regular step
                p = self.doc.add_paragraph()
                indent = 0.3 + indent_level * 0.15
                p.paragraph_format.left_indent = Cm(indent)
                p.paragraph_format.space_after = Pt(2)

                # Check for sub-numbering (e.g., 3.1., 3.2.)
                has_subnumber = False
                for i in range(len(stripped)):
                    if stripped[i].isdigit() or stripped[i] == '.':
                        continue
                    else:
                        if stripped[i] == ' ' and '.' in stripped[:i]:
                            has_subnumber = True
                        break

                if has_subnumber:
                    parts = stripped.split(' ', 1)
                    if len(parts) == 2:
                        run_num = p.add_run(parts[0] + " ")
                        run_num.bold = True
                        run_num.font.size = Pt(10)
                        run_num.font.color.rgb = DocColors.STEEL_BLUE
                        run_num.font.name = 'Calibri'

                        run_text = p.add_run(parts[1])
                        run_text.font.size = Pt(10)
                        run_text.font.name = 'Calibri'
                    else:
                        run = p.add_run(stripped)
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                else:
                    run = p.add_run(stripped)
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

    # ========================================================================
    # DIRECTIONAL SECTION
    # ========================================================================

    def _create_directional_section(self):
        """بخش حفاری جهتی"""
        self.section_counter += 1
        dp = self.project.directional_plan

        self._add_section_heading(
            f"{self.section_counter}. DIRECTIONAL DRILLING PLAN",
            level=1)

        # Summary
        self._add_section_heading("Directional Summary", level=2)

        dir_data = [
            ("Survey Tool", dp.survey_tool),
            ("Survey Interval", f"{dp.survey_frequency:.0f} ft"),
            ("Kickoff Point (MD)",
             f"{dp.kickoff_point_md:,.0f} ft"),
            ("Kickoff Point (TVD)",
             f"{dp.kickoff_point_tvd:,.0f} ft"),
            ("Build Rate", f"{dp.build_rate:.2f} °/100ft"),
            ("Turn Rate", f"{dp.turn_rate:.2f} °/100ft"),
            ("Hold Inclination", f"{dp.hold_inclination:.1f}°"),
            ("Hold Azimuth", f"{dp.hold_azimuth:.1f}°"),
            ("Target Inclination",
             f"{dp.target_inclination:.1f}°"),
            ("Target Azimuth", f"{dp.target_azimuth:.1f}°"),
            ("Maximum DLS", f"{dp.max_dls:.2f} °/100ft"),
            ("Horizontal Displacement",
             f"{dp.horizontal_displacement:,.0f} ft"),
            ("Vertical Section",
             f"{dp.vertical_section:,.0f} ft"),
        ]

        TableHelper.create_key_value_table(
            self.doc, dir_data, num_cols=4)

        # Wellpath Data
        if dp.wellpath_data:
            self._add_section_heading("Planned Wellpath", level=2)

            wp_headers = [
                "MD\n(ft)", "TVD\n(ft)", "Inc\n(°)",
                "Azi\n(°)", "DLS\n(°/100ft)",
                "N/S\n(ft)", "E/W\n(ft)", "VS\n(ft)"
            ]

            wp_data = []
            for wp in dp.wellpath_data:
                wp_data.append([
                    f"{wp.get('md', 0):,.0f}",
                    f"{wp.get('tvd', 0):,.0f}",
                    f"{wp.get('inclination', 0):.2f}",
                    f"{wp.get('azimuth', 0):.2f}",
                    f"{wp.get('dls', 0):.2f}",
                    str(wp.get('ns', '')),
                    str(wp.get('ew', '')),
                    str(wp.get('vs', '')),
                ])

            TableHelper.create_professional_table(
                self.doc, wp_headers, wp_data,
                caption=f"Table {self._next_table()}: "
                        f"Planned Wellpath Survey Data"
            )

        self.doc.add_page_break()

    # ========================================================================
    # BOP & WELL CONTROL SECTION
    # ========================================================================

    def _create_bop_well_control_section(self):
        """بخش BOP و کنترل چاه"""
        self.section_counter += 1
        bop = self.project.bop_stack
        wc = self.project.well_control

        self._add_section_heading(
            f"{self.section_counter}. BOP & WELL CONTROL", level=1)

        # BOP Configuration
        self._add_section_heading("BOP Stack Configuration", level=2)

        bop_data = [
            ("BOP Type", bop.bop_type),
            ("Working Pressure",
             f"{bop.working_pressure:,.0f} psi"),
            ("Bore Size", f'{bop.bore_size}"'),
            ("Manufacturer", bop.manufacturer),
            ("Model", bop.model),
            ("Annular Preventer",
             f'{bop.annular_preventer_size}" @ '
             f'{bop.annular_preventer_wp:,.0f} psi'),
            ("Pipe Ram Size", bop.pipe_ram_size),
            ("Blind Rams", "Yes" if bop.blind_ram else "No"),
            ("Shear Rams", "Yes" if bop.shear_ram else "No"),
            ("Variable Bore Rams",
             "Yes" if bop.variable_bore_ram else "No"),
            ("Kill Line",
             f'{bop.kill_line_size}" ID'),
            ("Choke Line",
             f'{bop.choke_line_size}" ID'),
            ("Choke Manifold WP",
             f"{bop.choke_manifold_wp:,.0f} psi"),
            ("Accumulator Capacity",
             f"{bop.accumulator_capacity:,.0f} gal"),
            ("Accumulator Precharge",
             f"{bop.accumulator_precharge:,.0f} psi"),
            ("Diverter Size",
             f'{bop.diverter_size}"'),
            ("Diverter Line",
             f'{bop.diverter_line_size}" ID'),
        ]

        TableHelper.create_key_value_table(
            self.doc, bop_data, num_cols=4)

        # BOP Testing
        self._add_section_heading("BOP Testing Requirements", level=2)

        test_data = [
            ("Function Test Frequency",
             bop.function_test_frequency),
            ("Pressure Test Frequency",
             bop.pressure_test_frequency),
            ("Low Pressure Test",
             f"{bop.bop_test_pressure_low:,.0f} psi - 5 min hold"),
            ("High Pressure Test",
             f"{bop.bop_test_pressure_high:,.0f} psi - 10 min hold"),
        ]

        TableHelper.create_key_value_table(self.doc, test_data)

        # Well Control Data
        self._add_section_heading("Well Control Data", level=2)

        wc_data = [
            ("Primary Kill Method", wc.kill_method),
            ("MAASP at Surface",
             f"{wc.maasp_surface:,.0f} psi"),
            ("Kick Tolerance", f"{wc.kick_tolerance:.0f} bbl"),
            ("Pit Gain Alarm Level",
             f"{wc.pit_gain_action_level:.0f} bbl"),
            ("Slow Pump Rate #1",
             f"{wc.slow_pump_rate_1:.0f} SPM @ "
             f"{wc.slow_pump_pressure_1:,.0f} psi"),
            ("Slow Pump Rate #2",
             f"{wc.slow_pump_rate_2:.0f} SPM @ "
             f"{wc.slow_pump_pressure_2:,.0f} psi"),
        ]

        TableHelper.create_key_value_table(
            self.doc, wc_data, num_cols=4)

        # Emergency contacts
        if wc.emergency_contacts:
            self._add_section_heading("Emergency Contacts", level=2)
            p = self.doc.add_paragraph(wc.emergency_contacts)
            for run in p.runs:
                run.font.size = Pt(10)

        self.doc.add_page_break()

    # ========================================================================
    # EVALUATION SECTION
    # ========================================================================

    def _create_evaluation_section(self):
        """بخش برنامه ارزیابی"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. EVALUATION PROGRAM", level=1)

        p = self.doc.add_paragraph()
        run = p.add_run(
            "The evaluation program includes wireline logging, "
            "LWD/MWD data acquisition, coring, formation testing, "
            "and cement evaluation. "
            "All evaluation activities shall be coordinated with "
            "the geology and reservoir engineering teams."
        )
        run.font.size = Pt(10)

        self._add_note(
            "Detailed evaluation program and tool selection "
            "to be finalized with service provider prior to operations.",
            "NOTE"
        )

        self.doc.add_page_break()

    # ========================================================================
    # TIME ESTIMATE SECTION
    # ========================================================================

    def _create_time_estimate_section(self):
        """بخش تخمین زمان"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. TIME vs DEPTH ESTIMATE", level=1)

        if not self.project.time_estimates:
            self.doc.add_paragraph(
                "No time estimate data entered.")
            self.doc.add_page_break()
            return

        # Time Summary Table
        headers = [
            "Section", "Operation",
            "Depth From\n(ft)", "Depth To\n(ft)",
            "Section\nDays", "Cumulative\nDays"
        ]

        data = []
        for te in self.project.time_estimates:
            data.append([
                te.section_name,
                te.operation,
                f"{te.depth_from:,.0f}" if te.depth_from > 0 else "-",
                f"{te.depth_to:,.0f}" if te.depth_to > 0 else "-",
                f"{te.total_section_days:.1f}",
                f"{te.cumulative_days:.1f}",
            ])

        TableHelper.create_professional_table(
            self.doc, headers, data,
            caption=f"Table {self._next_table()}: "
                    f"Time vs Depth Estimate (AFE)"
        )

        # Total
        if self.project.time_estimates:
            total_days = max(
                te.cumulative_days
                for te in self.project.time_estimates
            )
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                f"\nEstimated Total Well Duration: "
                f"{total_days:.1f} Days")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = DocColors.DARK_NAVY

        self._add_note(
            "Time estimate includes NPT contingency. "
            "Actual performance may vary based on hole conditions, "
            "weather, and operational factors.",
            "NOTE"
        )

        self.doc.add_page_break()

    # ========================================================================
    # ALL PROCEDURES SECTION (Part 3 continues to Part 4)
    # ========================================================================

    def _create_all_procedures_section(self):
        """بخش تمامی پروسیجرها"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. DETAILED OPERATING PROCEDURES",
            level=1)

        p = self.doc.add_paragraph()
        run = p.add_run(
            "This section contains detailed step-by-step procedures "
            "for all major drilling operations. "
            "All procedures are based on industry best practices, "
            "API standards, and company operating guidelines. "
            "These procedures shall be reviewed by all crew members "
            "before commencing each operation."
        )
        run.font.size = Pt(10)

        self.doc.add_paragraph("")

        # Generate all procedures
        procedures = self.proc_gen.generate_all_procedures()

        # Section numbering
        sub_section = 1

        # Pre-Spud
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Pre-Spud Checklist", level=2)
        self._add_procedure_steps(procedures.get('pre_spud', []))
        self.doc.add_page_break()
        sub_section += 1

        # Conductor
        if 'conductor' in procedures:
            self._add_section_heading(
                f"{self.section_counter}.{sub_section} "
                f"Conductor Setting Procedure", level=2)
            self._add_procedure_steps(procedures['conductor'])
            self.doc.add_page_break()
            sub_section += 1

        # Drilling Procedures per section
        for casing in self.project.casing_design:
            section_key = casing.section_name.lower().replace(' ', '_')
            drill_key = f'drill_{section_key}'

            if drill_key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{sub_section} "
                    f"Drilling Procedure - {casing.section_name} Section",
                    level=2)
                self._add_procedure_steps(procedures[drill_key])
                self.doc.add_page_break()
                sub_section += 1

        # Tripping Procedures
        for casing in self.project.casing_design:
            section_key = casing.section_name.lower().replace(' ', '_')
            trip_key = f'trip_{section_key}'

            if trip_key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{sub_section} "
                    f"Tripping Procedure - {casing.section_name} Section",
                    level=2)
                self._add_procedure_steps(procedures[trip_key])
                self.doc.add_page_break()
                sub_section += 1

        # Casing Running
        for casing in self.project.casing_design:
            section_key = casing.section_name.lower().replace(' ', '_')
            csg_key = f'casing_run_{section_key}'

            if csg_key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{sub_section} "
                    f"Casing Running - {casing.casing_od}\" "
                    f"{casing.section_name}", level=2)
                self._add_procedure_steps(procedures[csg_key])
                self.doc.add_page_break()
                sub_section += 1

        # Cementing
        for casing in self.project.casing_design:
            section_key = casing.section_name.lower().replace(' ', '_')
            cem_key = f'cement_{section_key}'

            if cem_key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{sub_section} "
                    f"Cementing Procedure - {casing.section_name}",
                    level=2)
                self._add_procedure_steps(procedures[cem_key])
                self.doc.add_page_break()
                sub_section += 1

        # BOP Procedures
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"BOP Nipple Up Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('bop_nipple_up', []))
        self.doc.add_page_break()
        sub_section += 1

        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"BOP Test Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('bop_test', []))
        self.doc.add_page_break()
        sub_section += 1

        # LOT/FIT
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"LOT / FIT Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('lot_fit', []))
        self.doc.add_page_break()
        sub_section += 1

        # Well Control
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Well Control Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('well_control', []))
        self.doc.add_page_break()
        sub_section += 1

        # Directional
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Directional Drilling Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('directional', []))
        self.doc.add_page_break()
        sub_section += 1

        # Logging
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Wireline Logging Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('logging', []))
        self.doc.add_page_break()
        sub_section += 1

        # Stuck Pipe
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Stuck Pipe Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('stuck_pipe', []))
        self.doc.add_page_break()
        sub_section += 1

        # Lost Circulation
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Lost Circulation Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('lost_circulation', []))
        self.doc.add_page_break()
        sub_section += 1

        # Wellbore Cleanup
        self._add_section_heading(
            f"{self.section_counter}.{sub_section} "
            f"Wellbore Cleanup Procedure", level=2)
        self._add_procedure_steps(
            procedures.get('wellbore_cleanup', []))
        self.doc.add_page_break()
        sub_section += 1

    # ========================================================================
    # EMERGENCY PROCEDURES
    # ========================================================================

    def _create_emergency_section(self):
        """بخش پروسیجرهای اضطراری"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. EMERGENCY PROCEDURES", level=1)

        procedures = self.proc_gen.generate_all_procedures()

        # H2S
        self._add_section_heading(
            f"{self.section_counter}.1 H₂S Emergency Procedure",
            level=2)
        self._add_procedure_steps(
            procedures.get('h2s_emergency', []))
        self.doc.add_page_break()

        # Kick Drill
        self._add_section_heading(
            f"{self.section_counter}.2 Well Control Drill Procedure",
            level=2)
        self._add_procedure_steps(
            procedures.get('kick_drill', []))
        self.doc.add_page_break()

        # Fishing
        self._add_section_heading(
            f"{self.section_counter}.3 Fishing Operations Procedure",
            level=2)
        self._add_procedure_steps(
            procedures.get('fishing', []))
        self.doc.add_page_break()

        # HSE
        self._add_section_heading(
            f"{self.section_counter}.4 HSE Requirements", level=2)
        self._add_procedure_steps(
            procedures.get('hse', []))
        self.doc.add_page_break()

        # Abandonment
        self._add_section_heading(
            f"{self.section_counter}.5 Well Abandonment Procedure",
            level=2)
        self._add_procedure_steps(
            procedures.get('abandonment', []))
        self.doc.add_page_break()

    # ========================================================================
    # APPENDICES
    # ========================================================================

    def _create_appendices(self):
        """بخش ضمایم"""
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. APPENDICES", level=1)

        appendices = [
            "Appendix A: Wellbore Schematic",
            "Appendix B: BHA Diagrams",
            "Appendix C: Casing Design Summary",
            "Appendix D: Directional Survey Program",
            "Appendix E: Cement Program Details",
            "Appendix F: Mud Program Details",
            "Appendix G: Hydraulics Calculations",
            "Appendix H: Torque & Drag Analysis",
            "Appendix I: Kill Sheet",
            "Appendix J: Emergency Contact List",
            "Appendix K: Rig Floor Poster",
            "Appendix L: Offset Well Data",
            "Appendix M: Material Safety Data Sheets (MSDS)",
            "Appendix N: Regulatory Permits",
        ]

        for app in appendices:
            p = self.doc.add_paragraph()
            run = p.add_run(app)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DocColors.DARK_NAVY
            run.font.name = 'Calibri'

            p2 = self.doc.add_paragraph()
            run2 = p2.add_run("  [To be attached separately]")
            run2.font.size = Pt(9)
            run2.font.italic = True
            run2.font.color.rgb = DocColors.GRAY

            self.doc.add_paragraph("")

        # End of Document
        self.doc.add_page_break()
        end_para = self.doc.add_paragraph()
        end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_para.paragraph_format.space_before = Pt(100)

        self._add_colored_line(DocColors.DARK_NAVY, width=3)

        end_run = end_para.add_run(
            "\n\n--- END OF DRILLING PROGRAM ---\n\n")
        end_run.bold = True
        end_run.font.size = Pt(16)
        end_run.font.color.rgb = DocColors.DARK_NAVY

        ci = self.project.company_info
        end_run2 = end_para.add_run(
            f"{ci.operator_name}\n"
            f"Well: {ci.well_name}\n"
            f"Document: {ci.document_number}\n"
            f"Revision: {ci.revision}\n"
            f"Date: {datetime.now().strftime('%d-%B-%Y')}\n"
        )
        end_run2.font.size = Pt(10)
        end_run2.font.color.rgb = DocColors.MEDIUM_GRAY

        self._add_colored_line(DocColors.DARK_NAVY, width=3)

    # ========================================================================
    # HEADERS AND FOOTERS
    # ========================================================================

    def _add_headers_footers(self):
        """اضافه کردن سرصفحه و پاصفحه"""
        ci = self.project.company_info

        for section in self.doc.sections:
            try:
                # ---- HEADER ----
                header = section.header
                header.is_linked_to_previous = False

                for para in header.paragraphs:
                    para.clear()

                if header.paragraphs:
                    hp = header.paragraphs[0]
                else:
                    hp = header.add_paragraph()

                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hp.clear()

                run_left = hp.add_run(ci.operator_name or "Operator")
                run_left.font.size = Pt(8)
                run_left.font.color.rgb = DocColors.NAVY
                run_left.font.bold = True

                run_sep1 = hp.add_run("  |  ")
                run_sep1.font.size = Pt(8)

                run_center = hp.add_run(
                    f"Drilling Program - {ci.well_name or 'Well'}")
                run_center.font.size = Pt(8)
                run_center.font.color.rgb = DocColors.DARK_NAVY
                run_center.font.bold = True

                run_sep2 = hp.add_run("  |  ")
                run_sep2.font.size = Pt(8)

                run_right = hp.add_run(
                    f"{ci.document_number or 'DRL-PRG-001'} "
                    f"Rev.{ci.revision or '0'}")
                run_right.font.size = Pt(8)
                run_right.font.color.rgb = DocColors.GRAY

                # Border under header
                try:
                    pPr = hp._p.get_or_add_pPr()
                    pBdr = parse_xml(
                        f'<w:pBdr {nsdecls("w")}>'
                        f'  <w:bottom w:val="single" w:sz="4" '
                        f'w:space="1" w:color="0F3460"/>'
                        f'</w:pBdr>'
                    )
                    pPr.append(pBdr)
                except Exception:
                    pass

                # ---- FOOTER ----
                footer = section.footer
                footer.is_linked_to_previous = False

                for para in footer.paragraphs:
                    para.clear()

                if footer.paragraphs:
                    fp = footer.paragraphs[0]
                else:
                    fp = footer.add_paragraph()

                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fp.clear()

                run_cls = fp.add_run(ci.classification.upper())
                run_cls.font.size = Pt(7)
                run_cls.font.color.rgb = DocColors.RED
                run_cls.font.bold = True

                run_sep3 = fp.add_run("  |  ")
                run_sep3.font.size = Pt(7)

                run_date = fp.add_run(
                    f"Printed: {datetime.now().strftime('%d-%b-%Y')}")
                run_date.font.size = Pt(7)
                run_date.font.color.rgb = DocColors.GRAY

                run_sep4 = fp.add_run("  |  Page ")
                run_sep4.font.size = Pt(7)
                run_sep4.font.color.rgb = DocColors.GRAY

                # Page number field
                try:
                    run_pg = fp.add_run()
                    run_pg._r.append(parse_xml(
                        f'<w:fldChar {nsdecls("w")} '
                        f'w:fldCharType="begin"/>'))
                    run_pg.font.size = Pt(7)

                    run_pg2 = fp.add_run()
                    run_pg2._r.append(parse_xml(
                        f'<w:instrText {nsdecls("w")} '
                        f'xml:space="preserve"> PAGE </w:instrText>'))
                    run_pg2.font.size = Pt(7)

                    run_pg3 = fp.add_run()
                    run_pg3._r.append(parse_xml(
                        f'<w:fldChar {nsdecls("w")} '
                        f'w:fldCharType="end"/>'))
                    run_pg3.font.size = Pt(7)

                except Exception:
                    pass

            except Exception as e:
                import logging
                logging.getLogger('DrillingProgram').warning(
                    f"Header/footer error: {e}")
                    
    def _add_single_header_footer(self, section, ci):
        """Helper - adds header and footer to one section"""
        # ---- HEADER ----
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing content
        for para in header.paragraphs:
            para.clear()

        # Use simple paragraph-based header (more compatible)
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()

        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.clear()

        # Left part
        run_left = hp.add_run(ci.operator_name or "Operator")
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = DocColors.NAVY
        run_left.font.bold = True

        # Separator
        run_sep = hp.add_run("  |  ")
        run_sep.font.size = Pt(8)
        run_sep.font.color.rgb = DocColors.GRAY

        # Center
        run_center = hp.add_run(
            f"Drilling Program - {ci.well_name or 'Well'}")
        run_center.font.size = Pt(8)
        run_center.font.color.rgb = DocColors.DARK_NAVY
        run_center.font.bold = True

        # Separator
        run_sep2 = hp.add_run("  |  ")
        run_sep2.font.size = Pt(8)
        run_sep2.font.color.rgb = DocColors.GRAY

        # Right
        run_right = hp.add_run(
            f"{ci.document_number or 'DRL-PRG-001'} Rev.{ci.revision or '0'}")
        run_right.font.size = Pt(8)
        run_right.font.color.rgb = DocColors.GRAY

        # Add border line under header
        try:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            pPr = hp._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>' +
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="0F3460"/>' +
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
        except Exception:
            pass

        # ---- FOOTER ----
        footer = section.footer
        footer.is_linked_to_previous = False

        for para in footer.paragraphs:
            para.clear()

        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()

        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()

        # Classification
        run_cls = fp.add_run(f"{ci.classification.upper()}")
        run_cls.font.size = Pt(7)
        run_cls.font.color.rgb = DocColors.RED
        run_cls.font.bold = True

        run_sep3 = fp.add_run("  |  ")
        run_sep3.font.size = Pt(7)
        run_sep3.font.color.rgb = DocColors.GRAY

        # Date
        run_date = fp.add_run(
            f"Printed: {datetime.now().strftime('%d-%b-%Y')}")
        run_date.font.size = Pt(7)
        run_date.font.color.rgb = DocColors.GRAY

        run_sep4 = fp.add_run("  |  Page ")
        run_sep4.font.size = Pt(7)
        run_sep4.font.color.rgb = DocColors.GRAY

        # Page number field
        try:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            run_pg = fp.add_run()
            fldChar1 = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run_pg._r.append(fldChar1)

            run_pg2 = fp.add_run()
            instrText = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run_pg2._r.append(instrText)

            run_pg3 = fp.add_run()
            fldChar2 = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run_pg3._r.append(fldChar2)

            for r in [run_pg, run_pg2, run_pg3]:
                r.font.size = Pt(7)
                r.font.color.rgb = DocColors.GRAY
        except Exception:
            run_pg_simple = fp.add_run("N/A")
            run_pg_simple.font.size = Pt(7)

